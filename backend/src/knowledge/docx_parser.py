"""
Patent docx parser — extracts structured content and formatting metadata
from finalized patent docx files in the 定稿文件/ directory.

File naming convention:
  A-*.docx = disclosure/conversation transcript (raw content for LLM reference)
  B-*.docx = finalized patent document (format + content reference)
"""

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm, Emu, Pt
from loguru import logger


# ─── Data Models ─────────────────────────────────────────────────────────────


@dataclass
class FormatMetadata:
    """Formatting metadata extracted from a reference patent docx."""

    # Page layout
    page_width_cm: float = 21.0
    page_height_cm: float = 29.7

    # Section margins (different sections may have different margins)
    # key = section purpose, value = (top, bottom, left, right) in cm
    section_margins: Dict[str, Tuple[float, float, float, float]] = field(
        default_factory=dict
    )

    # Body text formatting
    body_font_name: str = "楷体"
    body_font_size_pt: float = 14.0
    body_first_line_indent_cm: float = 0.99
    body_line_spacing_emu: int = 292100
    body_alignment: str = "JUSTIFY"

    # Section heading formatting (技术领域, 背景技术, etc.)
    section_heading_font_name: str = "楷体"
    section_heading_font_size_pt: float = 14.0
    section_heading_bold: bool = True

    # Document-level heading formatting (权利要求书, 说明书, etc.)
    doc_heading_font_name: str = "楷体"
    doc_heading_alignment: str = "CENTER"
    doc_heading_style: str = "Heading 1"


@dataclass
class PatentSection:
    """A section of a patent document."""

    name: str  # e.g. "技术领域", "背景技术", "发明内容"
    content: str  # full text content
    paragraphs: List[str] = field(default_factory=list)


@dataclass
class ParsedPatentDocx:
    """Structured content extracted from a B-file (finalized patent)."""

    title: str = ""
    ipc_code: str = ""
    source_file: str = ""

    # Document sections
    abstract: str = ""
    claims: List[str] = field(default_factory=list)  # each claim as a string
    description_sections: Dict[str, PatentSection] = field(default_factory=dict)

    # Full text for LLM reference
    full_text: str = ""

    # Formatting metadata
    format_metadata: Optional[FormatMetadata] = None


@dataclass
class ParsedDisclosure:
    """Content extracted from an A-file (disclosure/conversation)."""

    title: str = ""
    source_file: str = ""
    full_text: str = ""


@dataclass
class FinalizedPatentEntry:
    """A complete finalized patent entry with both A and B file data."""

    directory_name: str
    patent_doc: Optional[ParsedPatentDocx] = None  # B-file
    disclosure_doc: Optional[ParsedDisclosure] = None  # A-file


# ─── Document Section Markers ────────────────────────────────────────────────

# Heading markers used to split the document into major sections
_DOC_HEADINGS = [
    "说明书摘要",
    "说   明   书   摘   要",
    "摘要附图",
    "摘   要   附   图",
    "权利要求书",
    "权    利    要    求    书",
    "说明书",
    "说    明    书",
    "说明书附图",
    "说   明   书   附   图",
]

# Sub-section headings within 说明书
_DESC_SECTION_HEADINGS = [
    "技术领域",
    "背景技术",
    "发明内容",
    "附图说明",
    "具体实施方式",
]


def _normalize_heading(text: str) -> str:
    """Remove extra spaces from spaced-out headings."""
    return re.sub(r"\s+", "", text.strip())


def _is_doc_heading(text: str) -> Optional[str]:
    """Check if text is a document-level heading, return normalized name."""
    normalized = _normalize_heading(text)
    mapping = {
        "说明书摘要": "摘要",
        "摘要附图": "摘要附图",
        "权利要求书": "权利要求书",
        "说明书": "说明书",
        "说明书附图": "说明书附图",
    }
    return mapping.get(normalized)


def _is_section_heading(text: str, bold: bool) -> Optional[str]:
    """Check if text is a description sub-section heading."""
    stripped = text.strip()
    if stripped in _DESC_SECTION_HEADINGS and bold:
        return stripped
    return None


# ─── Format Extraction ───────────────────────────────────────────────────────


def extract_format_metadata(doc: Document) -> FormatMetadata:
    """Extract formatting metadata from a reference patent docx."""
    meta = FormatMetadata()

    # Page size from first section
    if doc.sections:
        first_section = doc.sections[0]
        meta.page_width_cm = round(first_section.page_width.cm, 2)
        meta.page_height_cm = round(first_section.page_height.cm, 2)

    # Section margins
    section_names = ["摘要", "摘要附图", "权利要求书", "说明书", "说明书附图"]
    for i, section in enumerate(doc.sections):
        name = section_names[i] if i < len(section_names) else f"section_{i}"
        meta.section_margins[name] = (
            round(section.top_margin.cm, 2),
            round(section.bottom_margin.cm, 2),
            round(section.left_margin.cm, 2),
            round(section.right_margin.cm, 2),
        )

    # Body text formatting — sample from paragraphs with "Body Text First Indent" style
    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ""
        if "Body Text" in style_name and para.runs and para.text.strip():
            run = para.runs[0]
            if run.font.name:
                meta.body_font_name = run.font.name
            if run.font.size:
                meta.body_font_size_pt = run.font.size.pt

            fmt = para.paragraph_format
            if fmt.first_line_indent:
                meta.body_first_line_indent_cm = round(fmt.first_line_indent.cm, 2)
            if fmt.line_spacing:
                if isinstance(fmt.line_spacing, int):
                    meta.body_line_spacing_emu = fmt.line_spacing
                else:
                    # It's a float ratio
                    meta.body_line_spacing_emu = int(fmt.line_spacing * 240 * 914.4)
            break

    # Section heading formatting — find bold paragraphs matching known headings
    for para in doc.paragraphs:
        text = para.text.strip()
        if text in _DESC_SECTION_HEADINGS and para.runs:
            run = para.runs[0]
            if run.font.bold:
                if run.font.name:
                    meta.section_heading_font_name = run.font.name
                if run.font.size:
                    meta.section_heading_font_size_pt = run.font.size.pt
                meta.section_heading_bold = True
                break

    return meta


# ─── Content Parsing ─────────────────────────────────────────────────────────


def parse_patent_docx(file_path: Path) -> ParsedPatentDocx:
    """
    Parse a B-file (finalized patent docx) and extract structured content.

    Returns a ParsedPatentDocx with:
    - abstract text
    - claims list
    - description sections (技术领域, 背景技术, 发明内容, 附图说明, 具体实施方式)
    - full text for LLM reference
    - format metadata
    """
    doc = Document(str(file_path))
    result = ParsedPatentDocx(source_file=str(file_path))

    # Extract IPC code from filename
    filename = file_path.stem
    ipc_match = re.search(r"[A-H]\d{2}[A-Z]\d+", filename)
    if ipc_match:
        result.ipc_code = ipc_match.group()

    # Extract title from directory name
    result.title = file_path.parent.name

    # Extract format metadata
    result.format_metadata = extract_format_metadata(doc)

    # Split document into major sections based on Heading 1 paragraphs
    current_section: Optional[str] = None
    section_paragraphs: Dict[str, List[str]] = {
        "摘要": [],
        "摘要附图": [],
        "权利要求书": [],
        "说明书": [],
        "说明书附图": [],
    }

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Check if this is a document-level heading
        style_name = para.style.name if para.style else ""
        if "Heading" in style_name:
            heading_name = _is_doc_heading(text)
            if heading_name:
                current_section = heading_name
                continue

        # Accumulate paragraphs under current section
        if current_section and current_section in section_paragraphs:
            section_paragraphs[current_section].append(text)

    # Extract abstract
    result.abstract = "\n".join(section_paragraphs.get("摘要", []))

    # Extract claims
    claims_text = section_paragraphs.get("权利要求书", [])
    result.claims = _parse_claims_from_paragraphs(claims_text)

    # Extract description sub-sections
    desc_paragraphs = section_paragraphs.get("说明书", [])
    result.description_sections = _parse_description_sections(desc_paragraphs, doc)

    # Build full text
    all_text_parts = []
    if result.abstract:
        all_text_parts.append(f"【摘要】\n{result.abstract}")
    if result.claims:
        all_text_parts.append("【权利要求书】\n" + "\n".join(result.claims))
    for name, section in result.description_sections.items():
        all_text_parts.append(f"【{name}】\n{section.content}")
    result.full_text = "\n\n".join(all_text_parts)

    return result


def _parse_claims_from_paragraphs(paragraphs: List[str]) -> List[str]:
    """Parse individual claims from a flat list of paragraphs."""
    claims: List[str] = []
    current_claim_parts: List[str] = []

    claim_start_pattern = re.compile(r"^\d+[、．.]")

    for para in paragraphs:
        if claim_start_pattern.match(para):
            # Start of a new claim
            if current_claim_parts:
                claims.append("\n".join(current_claim_parts))
            current_claim_parts = [para]
        else:
            # Continuation of current claim
            current_claim_parts.append(para)

    # Don't forget the last claim
    if current_claim_parts:
        claims.append("\n".join(current_claim_parts))

    return claims


def _parse_description_sections(
    paragraphs: List[str], doc: Document
) -> Dict[str, PatentSection]:
    """Parse description into sub-sections based on bold headings."""
    sections: Dict[str, PatentSection] = {}
    current_heading: Optional[str] = None
    current_paragraphs: List[str] = []

    # Build a set of bold paragraph texts for quick lookup
    bold_texts = set()
    for para in doc.paragraphs:
        if para.text.strip() in _DESC_SECTION_HEADINGS:
            if para.runs and para.runs[0].font.bold:
                bold_texts.add(para.text.strip())

    for para in paragraphs:
        stripped = para.strip()
        if stripped in bold_texts:
            # Save previous section
            if current_heading and current_paragraphs:
                content = "\n".join(current_paragraphs)
                sections[current_heading] = PatentSection(
                    name=current_heading,
                    content=content,
                    paragraphs=current_paragraphs.copy(),
                )
            current_heading = stripped
            current_paragraphs = []
        else:
            if current_heading:
                current_paragraphs.append(para)

    # Save last section
    if current_heading and current_paragraphs:
        content = "\n".join(current_paragraphs)
        sections[current_heading] = PatentSection(
            name=current_heading,
            content=content,
            paragraphs=current_paragraphs.copy(),
        )

    return sections


def parse_disclosure_docx(file_path: Path) -> ParsedDisclosure:
    """
    Parse an A-file (disclosure/conversation transcript).
    Extracts all text content for LLM reference.
    """
    doc = Document(str(file_path))
    result = ParsedDisclosure(
        source_file=str(file_path),
        title=file_path.parent.name,
    )

    paragraphs = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
    result.full_text = "\n".join(paragraphs)

    return result


# ─── Directory Scanner ───────────────────────────────────────────────────────


def scan_finalized_patents_dir(base_dir: Path) -> List[FinalizedPatentEntry]:
    """
    Scan the 定稿文件/ directory and parse all patent entries.

    Expected structure:
        base_dir/
        ├── 专利名称1/
        │   ├── A-专利名称1.docx  (disclosure)
        │   └── B-专利名称1 IPC_CODE.docx  (finalized patent)
        ├── 专利名称2/
        │   ├── A-*.docx
        │   └── B-*.docx
        └── ...

    Returns a list of FinalizedPatentEntry objects.
    """
    if not base_dir.exists():
        logger.warning(f"定稿文件目录不存在: {base_dir}")
        return []

    entries: List[FinalizedPatentEntry] = []

    for subdir in sorted(base_dir.iterdir()):
        if not subdir.is_dir():
            continue

        entry = FinalizedPatentEntry(directory_name=subdir.name)

        # Find A-file and B-file
        for docx_file in subdir.glob("*.docx"):
            filename = docx_file.name
            if filename.startswith("A-"):
                try:
                    entry.disclosure_doc = parse_disclosure_docx(docx_file)
                    logger.debug(f"解析 A 文件: {docx_file.name}")
                except Exception as e:
                    logger.error(f"解析 A 文件失败 {docx_file}: {e}")
            elif filename.startswith("B-"):
                try:
                    entry.patent_doc = parse_patent_docx(docx_file)
                    logger.debug(f"解析 B 文件: {docx_file.name}")
                except Exception as e:
                    logger.error(f"解析 B 文件失败 {docx_file}: {e}")

        if entry.patent_doc or entry.disclosure_doc:
            entries.append(entry)
            logger.info(
                f"加载定稿专利: {subdir.name} "
                f"(B={'有' if entry.patent_doc else '无'}, "
                f"A={'有' if entry.disclosure_doc else '无'})"
            )

    logger.info(f"共加载 {len(entries)} 篇定稿专利")
    return entries


def get_reference_format(base_dir: Path) -> FormatMetadata:
    """
    Extract formatting metadata from the first available B-file.
    Used as the reference format for generating new patent documents.
    """
    if not base_dir.exists():
        logger.warning(f"定稿文件目录不存在，使用默认格式: {base_dir}")
        return FormatMetadata()

    for subdir in sorted(base_dir.iterdir()):
        if not subdir.is_dir():
            continue
        for docx_file in subdir.glob("B-*.docx"):
            try:
                doc = Document(str(docx_file))
                meta = extract_format_metadata(doc)
                logger.info(f"从 {docx_file.name} 提取参考格式")
                return meta
            except Exception as e:
                logger.error(f"提取格式失败 {docx_file}: {e}")
                continue

    logger.warning("未找到可用的参考格式文件，使用默认格式")
    return FormatMetadata()
