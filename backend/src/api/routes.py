from typing import Any, Dict, List
import asyncio
import uuid
from datetime import datetime

from fastapi import APIRouter, HTTPException, BackgroundTasks, Query, status
from fastapi.responses import FileResponse, StreamingResponse
from loguru import logger

from .schemas import (
    ChatMessageRequest,
    CreateTaskRequest,
    TaskResponse,
    TaskDetailResponse,
    TaskListResponse,
    WorkflowEventResponse,
    WorkflowListResponse,
    WorkflowResponse,
    SearchPatentRequest,
    SearchResponse,
    SystemStatusResponse,
    KnowledgeBaseSearchResponse,
    WorkflowStartRequest,
    AgentDetailResponse,
    AgentListResponse,
    AgentToggleResponse,
    AgentUpdateResponse,
    OrganizationUpdateResponse,
    OrgNodeResponse,
    PriorArtReferenceResponse,
    # Conversation schemas
    ConversationSummary,
    ConversationMessage,
    ConversationDetail,
    ConversationListResponse,
    ConversationChatResponse,
    CreateConversationRequest,
    ConversationChatRequest,
    CreateWorkflowFromConversationRequest,
)
from ..models.domain import PatentTask
from ..models.enums import WorkflowState
from ..core.workflow_engine import PatentWorkflowEngine, WorkflowContext
from ..knowledge.base import get_knowledge_base
from ..data_sources.base import get_data_source_manager
from ..infrastructure.persistence import get_store
from .schemas import WorkflowEventResponse, OrgNodeResponse


router = APIRouter(tags=["patent-agents"])


# 内存存储 - 生产环境请替换为数据库
tasks_store: Dict[str, PatentTask] = {}
task_events: Dict[str, List[WorkflowEventResponse]] = {}
workflow_lock = asyncio.Lock()

# 对话会话存储
conversations_store: Dict[str, dict] = {}
conversations_lock = asyncio.Lock()

workflow_engine = PatentWorkflowEngine()
organization_tree_store: OrgNodeResponse | None = None

# ── Agent Override Store (真实持久化) ──
from ..core.override_store import get_override_store

_override_store = get_override_store()

# ── Hermes Agent Service (真实 hermes-agent 集成) ──
from ..agents.hermes_agent_service import get_hermes_agent_service

_hermes_service = get_hermes_agent_service()

# ── 订阅Agent事件总线，存入task_events供SSE回放 ──
from ..core.events import (
    EventType,
    subscribe_event,
    BaseEvent,
    AgentThinkingEvent,
    AgentToolCallStartEvent,
    AgentToolCallEndEvent,
    AgentDispatchEvent,
    AgentContentEvent,
)


def _on_agent_event(event: BaseEvent):
    """将Agent事件存入task_events供/stream端点回放（备份路径）
    
    当workflow通过event_callback直接写入时，此处会产生重复。
    通过检查最近事件避免重复。
    """
    task_id = getattr(event, "task_id", None)
    if not task_id:
        return

    # 如果task_events中最后一条事件时间戳在1秒内且类型相同，跳过（已由callback写入）
    existing = task_events.get(task_id, [])
    if existing:
        last = existing[-1]
        if last.event_type == event.event_type.value:
            last_ts = last.timestamp
            now = datetime.now()
            if hasattr(last_ts, 'timestamp'):
                diff = (now - last_ts).total_seconds()
            else:
                diff = 0
            if diff < 2:
                return  # 跳过重复

    agent_name = getattr(event, "agent_name", "") or getattr(event, "from_agent", "Agent")
    event_type_str = event.event_type.value

    if isinstance(event, AgentThinkingEvent):
        message = f"💭 {event.thought[:200]}"
    elif isinstance(event, AgentToolCallStartEvent):
        message = f"🔧 调用工具: {event.tool_name}"
    elif isinstance(event, AgentToolCallEndEvent):
        status_icon = "✅" if event.success else "❌"
        message = f"{status_icon} {event.tool_name} 返回: {event.result[:150]}"
    elif isinstance(event, AgentDispatchEvent):
        agent_name = event.from_agent
        message = f"🎯 调度 → {event.to_agent}: {event.task_description[:100]}"
    elif isinstance(event, AgentContentEvent):
        message = f"📄 输出: {event.content[:200]}"
    else:
        message = "事件"

    logger.info(f"[AgentEvent] {event_type_str} | {agent_name} | {message[:80]}")

    task_events.setdefault(task_id, []).append(
        WorkflowEventResponse(
            task_id=task_id,
            timestamp=datetime.now(),
            agent=agent_name,
            message=message,
            event_type=event_type_str,
            data=event.to_dict(),
        )
    )


subscribe_event(EventType.AGENT_THINKING, _on_agent_event)
subscribe_event(EventType.AGENT_TOOL_CALL_START, _on_agent_event)
subscribe_event(EventType.AGENT_TOOL_CALL_END, _on_agent_event)
subscribe_event(EventType.AGENT_DISPATCH, _on_agent_event)
subscribe_event(EventType.AGENT_CONTENT, _on_agent_event)
# NOTE: 上述订阅作为备份路径（无event_callback时的fallback）
# 当workflow传入event_callback时，事件会通过callback直接写入task_events
# 为避免重复，_on_agent_event只在task_events中无该task的最近同类事件时才写入

# ── 持久化存储辅助 ──
_store_instance = None


def _get_persist_store():
    global _store_instance
    if _store_instance is None:
        _store_instance = get_store()
    return _store_instance


async def _persist_task(task_id: str) -> None:
    task = tasks_store.get(task_id)
    if task is None:
        return
    try:
        await _get_persist_store().save("tasks", task_id, task.model_dump(mode="json"))
    except Exception as e:
        logger.warning(f"保存任务 {task_id} 到数据库失败: {e}")


async def _persist_events(task_id: str) -> None:
    events = task_events.get(task_id)
    if events is None:
        return
    try:
        await _get_persist_store().save("task_events", task_id, [e.model_dump(mode="json") for e in events])
    except Exception as e:
        logger.warning(f"保存事件 {task_id} 到数据库失败: {e}")


async def _persist_conversation(conv_id: str) -> None:
    conv = conversations_store.get(conv_id)
    if conv is None:
        return
    try:
        await _get_persist_store().save("conversations", conv_id, conv)
    except Exception as e:
        logger.warning(f"保存对话 {conv_id} 到数据库失败: {e}")


async def _persist_workflow(task_id: str) -> None:
    """持久化WorkflowContext到数据库"""
    context = workflow_engine.get_workflow(task_id)
    if context is None:
        return
    try:
        data = _workflow_context_to_response(context).model_dump(mode="json")
        await _get_persist_store().save("workflows", task_id, data)
    except Exception as e:
        logger.warning(f"保存工作流 {task_id} 到数据库失败: {e}")


async def _persist_org_tree() -> None:
    if organization_tree_store is None:
        return
    try:
        await _get_persist_store().save("org_tree", "root", organization_tree_store.model_dump(mode="json"))
    except Exception as e:
        logger.warning(f"保存组织架构到数据库失败: {e}")


async def restore_stores_from_db() -> None:
    """启动时从数据库恢复内存存储"""
    store = _get_persist_store()
    restored_tasks = 0
    restored_events = 0
    restored_conversations = 0

    # 恢复任务
    try:
        for key, value in await store.load_all("tasks"):
            try:
                tasks_store[key] = PatentTask.model_validate(value)
                restored_tasks += 1
            except Exception as e:
                logger.warning(f"恢复任务 {key} 失败: {e}")
    except Exception as e:
        logger.warning(f"恢复任务列表失败: {e}")

    # 恢复事件
    try:
        for key, value in await store.load_all("task_events"):
            try:
                task_events[key] = [WorkflowEventResponse.model_validate(e) for e in value]
                restored_events += 1
            except Exception as e:
                logger.warning(f"恢复事件 {key} 失败: {e}")
    except Exception as e:
        logger.warning(f"恢复事件列表失败: {e}")

    # 恢复对话
    try:
        for key, value in await store.load_all("conversations"):
            conversations_store[key] = value
            restored_conversations += 1
    except Exception as e:
        logger.warning(f"恢复对话列表失败: {e}")

    # 恢复工作流
    restored_workflows = 0
    try:
        for key, value in await store.load_all("workflows"):
            try:
                # 通过workflow_engine重建WorkflowContext
                existing = workflow_engine.get_workflow(key)
                if not existing:
                    ctx = workflow_engine.create_workflow(
                        task_id=key,
                        user_id=value.get("user_id", "unknown"),
                        description="",  # 从持久化数据恢复
                    )
                    # 恢复状态
                    from src.core.workflow_engine import WorkflowState
                    ctx.current_phase = WorkflowState(value.get("current_state", "initialized"))
                    ctx.iteration_count = value.get("iteration_count", 0)
                    # 恢复输出
                    outputs = value.get("outputs", {})
                    ctx.requirement_analysis = outputs.get("requirement_analysis", {})
                    ctx.retrieval_report = outputs.get("retrieval_report", {})
                    ctx.patent_draft = outputs.get("patent_draft", {})
                    ctx.review_report = outputs.get("review_report", {})
                    ctx.brainstorming_output = outputs.get("brainstorming", {})
                    restored_workflows += 1
            except Exception as e:
                logger.warning(f"恢复工作流 {key} 失败: {e}")
    except Exception as e:
        logger.warning(f"恢复工作流列表失败: {e}")

    # 恢复组织架构
    try:
        org_val = await store.load("org_tree", "root")
        if org_val is not None:
            global organization_tree_store
            organization_tree_store = OrgNodeResponse.model_validate(org_val)
    except Exception as e:
        logger.warning(f"恢复组织架构失败: {e}")

    logger.info(
        f"从数据库恢复: {restored_tasks} 个任务, "
        f"{restored_events} 组事件, "
        f"{restored_conversations} 个对话, "
        f"{restored_workflows} 个工作流"
    )


def _workflow_context_to_response(context: WorkflowContext) -> WorkflowResponse:
    return WorkflowResponse(
        task_id=context.task_id,
        user_id=context.user_id,
        current_state=context.current_phase.value,
        created_at=context.created_at,
        updated_at=context.updated_at,
        iteration_count=context.iteration_count,
        message_count=len(context.message_history),
        phase_history=[
            {
                "phase": result.phase.value,
                "success": result.success,
                "duration_seconds": result.duration_seconds,
                "issues": result.issues,
                "warnings": result.warnings,
            }
            for result in context.phase_history
        ],
        outputs={
            "brainstorming": context.brainstorming_output,
            "requirement_analysis": context.requirement_analysis,
            "retrieval_report": context.retrieval_report,
            "patent_draft": context.patent_draft,
            "review_report": context.review_report,
        },
    )


def _get_agent_memory_stats(profile_id: str, dir_name: str) -> Dict[str, Any]:
    """从 hermes session 文件读取 Agent 的实际记忆统计数据和条目"""
    import os as _mem_os
    import json as _mem_json
    from pathlib import Path as _MemPath

    hermes_home = _MemPath(__file__).parent.parent.parent / "hermes_home"
    profile_sessions_dir = hermes_home / "profiles" / dir_name / "sessions"
    global_sessions_dir = hermes_home / "sessions"

    total_messages = 0
    total_sessions = 0
    total_size = 0
    last_updated = None
    entries: List[Dict[str, Any]] = []  # 记忆条目

    # 1. Profile-specific sessions (short_term memory)
    if profile_sessions_dir.is_dir():
        for f in sorted(profile_sessions_dir.glob("*.json"), key=lambda x: x.stat().st_mtime, reverse=True):
            try:
                with open(f, "r", encoding="utf-8") as fh:
                    data = _mem_json.load(fh)
                if isinstance(data, dict):
                    msg_count = data.get("message_count", 0)
                    total_messages += msg_count
                    total_sessions += 1
                    total_size += f.stat().st_size
                    updated = data.get("last_updated")
                    if updated and (last_updated is None or updated > last_updated):
                        last_updated = updated
                    # 提取消息作为记忆条目
                    for msg in data.get("messages", [])[-20:]:  # 最近20条
                        role = msg.get("role", "unknown")
                        content = msg.get("content", "")
                        if content and role in ("user", "assistant"):
                            ts = msg.get("timestamp") or updated or ""
                            entries.append({
                                "id": f"{f.stem}_{len(entries)}",
                                "type": "context" if role == "assistant" else "event",
                                "key": f"[{role}] {content[:50]}",
                                "value": content[:300],
                                "score": None,
                                "created_at": ts,
                                "updated_at": ts,
                                "tags": [role, data.get("session_id", "")[:20]],
                            })
            except Exception:
                pass

    # 2. Global sessions that match this profile
    if global_sessions_dir.is_dir():
        for f in sorted(global_sessions_dir.glob("*.json"), key=lambda x: x.stat().st_mtime, reverse=True):
            try:
                with open(f, "r", encoding="utf-8") as fh:
                    data = _mem_json.load(fh)
                if isinstance(data, dict):
                    session_id = data.get("session_id", "")
                    sys_prompt = data.get("system_prompt", "")
                    if dir_name in session_id or dir_name in sys_prompt:
                        msg_count = data.get("message_count", 0)
                        total_messages += msg_count
                        total_sessions += 1
                        total_size += f.stat().st_size
                        updated = data.get("last_updated")
                        if updated and (last_updated is None or updated > last_updated):
                            last_updated = updated
                        # 提取消息
                        for msg in data.get("messages", [])[-10:]:
                            role = msg.get("role", "unknown")
                            content = msg.get("content", "")
                            if content and role in ("user", "assistant"):
                                ts = msg.get("timestamp") or updated or ""
                                entries.append({
                                    "id": f"{f.stem}_{len(entries)}",
                                    "type": "context" if role == "assistant" else "event",
                                    "key": f"[{role}] {content[:50]}",
                                    "value": content[:300],
                                    "score": None,
                                    "created_at": ts,
                                    "updated_at": ts,
                                    "tags": [role, session_id[:20]],
                                })
            except Exception:
                pass

    # 限制总条目数（前端性能）
    entries = entries[:50]

    return {
        "short_term_count": total_messages,
        "short_term_size": total_size,
        "short_term_entries": entries,
        "long_term_count": total_sessions,
        "long_term_size": total_size,
        "long_term_entries": [
            {
                "id": f"session_{i}",
                "type": "fact",
                "key": f"会话 {i+1}",
                "value": e["key"][:80],
                "score": None,
                "created_at": e.get("created_at", ""),
                "updated_at": e.get("updated_at", ""),
                "tags": ["session"],
            }
            for i, e in enumerate(entries[:total_sessions])
        ] if entries else [],
        "knowledge_base_count": total_sessions,
        "knowledge_base_size": total_size,
        "knowledge_base_entries": [],
        "last_updated": last_updated,
    }


def _hermes_role_to_ui(role_str: str) -> str:
    """将 hermes config 中的 role 字符串映射为前端 UI role"""
    mapping = {
        "orchestrator": "orchestrator",
        "ceo": "orchestrator",
        "specialist": "specialist",
        "requirement_analyst": "specialist",
        "retrieval_analyst": "specialist",
        "patent_writer": "specialist",
        "assistant": "assistant",
        "brainstorm_partner": "assistant",
        "critic": "critic",
        "quality_reviewer": "critic",
    }
    return mapping.get(role_str, "specialist")


def _agent_tool_category(tool_name: str) -> str:
    if any(keyword in tool_name for keyword in ("search", "retrieval", "knowledge")):
        return "search"
    if any(keyword in tool_name for keyword in ("format", "write", "draft", "document", "claim")):
        return "file"
    if any(keyword in tool_name for keyword in ("delegate", "spawn", "workflow")):
        return "external"
    return "analysis"


def _require_agent_profile(agent_id: str):
    """验证 agent 存在于 Hermes Agent Service"""
    cfg = _hermes_service.get_config(agent_id)
    if not cfg:
        raise HTTPException(status_code=404, detail="Agent不存在")
    return cfg


def _validate_org_tree(tree: OrgNodeResponse, depth: int = 0, seen_ids: set[str] | None = None) -> int:
    if seen_ids is None:
        seen_ids = set()
    if depth > 8:
        raise HTTPException(status_code=400, detail="组织架构层级过深")
    if tree.id in seen_ids:
        raise HTTPException(status_code=400, detail="组织架构节点ID重复")
    if not tree.id.strip() or not tree.name.strip():
        raise HTTPException(status_code=400, detail="组织架构节点ID和名称不能为空")
    if tree.type == "agent" and tree.children:
        raise HTTPException(status_code=400, detail="Agent节点不能包含子节点")

    seen_ids.add(tree.id)
    node_count = 1
    for child in tree.children:
        node_count += _validate_org_tree(child, depth + 1, seen_ids)
        if node_count > 100:
            raise HTTPException(status_code=400, detail="组织架构节点数量过多")
    return node_count


def _profile_organization_tree() -> Dict[str, Any]:
    """基于 Hermes Agent Service 配置生成组织架构树"""
    all_configs = _hermes_service.get_all_configs()

    # 按 role 分组
    ceo_nodes = []
    analysis_nodes = []
    writing_nodes = []

    for cfg in all_configs:
        node = {
            "id": cfg.profile_id,
            "name": cfg.name,
            "type": "agent",
            "description": cfg.description,
            "expanded": True,
            "children": [],
        }
        ui_role = _hermes_role_to_ui(cfg.role)
        if ui_role == "orchestrator":
            ceo_nodes.append(node)
        elif ui_role == "critic":
            writing_nodes.append(node)
        elif cfg.role in ("patent_writer",):
            writing_nodes.append(node)
        else:
            analysis_nodes.append(node)

    return {
        "id": "root",
        "name": "专利智能体系统",
        "type": "team",
        "description": "基于 Hermes Agent Profiles 的多智能体组织架构",
        "expanded": True,
        "children": [
            {
                "id": "orchestration-group",
                "name": "统筹管理层",
                "type": "group",
                "description": "负责流程调度、质量门控与跨 Agent 协同",
                "expanded": True,
                "children": ceo_nodes,
            },
            {
                "id": "analysis-group",
                "name": "分析与头脑风暴层",
                "type": "group",
                "description": "负责前期对话澄清、需求分析与专利检索",
                "expanded": True,
                "children": analysis_nodes,
            },
            {
                "id": "writing-group",
                "name": "撰写与审查层",
                "type": "group",
                "description": "负责专利申请文件生成与质量审查",
                "expanded": True,
                "children": writing_nodes,
            },
        ],
    }


def _append_workflow_event(task_id: str, agent: str, message: str, event_type: str, data=None):
    task_events.setdefault(task_id, []).append(
        WorkflowEventResponse(
            task_id=task_id,
            timestamp=datetime.now(),
            agent=agent,
            message=message,
            event_type=event_type,
            data=data or {},
        )
    )


@router.post("/workflows", status_code=status.HTTP_201_CREATED)
async def create_workflow_session(request: WorkflowStartRequest):
    """创建 Hermes/Profile 驱动的专利工作流会话（默认先进入头脑风暴）"""
    task_id = request.task_id or str(uuid.uuid4())
    async with workflow_lock:
        context = workflow_engine.create_workflow(
            task_id=task_id,
            user_id=request.user_id,
            description=request.tech_description,
            patent_type_preference=(
                request.patent_type_preference.value
                if request.patent_type_preference is not None
                else None
            ),
        )
        task_events[task_id] = []
        _append_workflow_event(
            task_id=task_id,
            agent="workflow_engine",
            message="专利工作流会话已创建，进入头脑风暴阶段",
            event_type="workflow.created",
            data={"state": context.current_phase.value},
        )
    await _persist_events(task_id)
    await _persist_workflow(task_id)
    return _workflow_context_to_response(context)


@router.post("/workflows/{task_id}/chat")
async def chat_with_brainstorm_agent(task_id: str, request: ChatMessageRequest):
    """与头脑风暴 Agent 继续讨论专利细节"""
    try:
        response = await workflow_engine.add_chat_message(
            task_id=task_id,
            role="user",
            content=request.content,
        )
    except ValueError as e:
        raise HTTPException(status_code=404, detail=str(e)) from e

    async with workflow_lock:
        _append_workflow_event(
            task_id=task_id,
            agent="brainstorm_partner",
            message=response.get("content", "消息已记录"),
            event_type="chat.message.created",
            data=response,
        )
    await _persist_events(task_id)
    return response


@router.post("/workflows/{task_id}/start")
async def start_workflow(task_id: str, background_tasks: BackgroundTasks):
    """确认专利细节后，启动完整专利申请流程"""
    async with workflow_lock:
        context = workflow_engine.get_workflow(task_id)
    if not context:
        raise HTTPException(status_code=404, detail="工作流不存在")

    async def phase_callback(phase, result):
        async with workflow_lock:
            _append_workflow_event(
                task_id=task_id,
                agent=phase.value,
                message=f"阶段 {phase.value} 已完成",
                event_type="workflow.phase.completed",
                data={
                    "phase": phase.value,
                    "success": result.success,
                    "duration_seconds": result.duration_seconds,
                    "issues": result.issues,
                },
            )

    def _workflow_event_callback(agent_name: str, event_type: str, message: str, data: Dict[str, Any]):
        """直接将agent事件写入task_events（绕过事件总线）"""
        task_events.setdefault(task_id, []).append(
            WorkflowEventResponse(
                task_id=task_id,
                timestamp=datetime.now(),
                agent=agent_name,
                message=message,
                event_type=event_type,
                data=data,
            )
        )

    async def run_workflow():
        try:
            await workflow_engine.execute_full_workflow(
                context,
                phase_callback=phase_callback,
                event_callback=_workflow_event_callback,
            )
            async with workflow_lock:
                _append_workflow_event(
                    task_id=task_id,
                    agent="workflow_engine",
                    message="专利申请流程已完成",
                    event_type="workflow.completed",
                    data={"state": context.current_phase.value},
                )
            await _persist_events(task_id)
            await _persist_workflow(task_id)
        except Exception as e:
            async with workflow_lock:
                _append_workflow_event(
                    task_id=task_id,
                    agent="workflow_engine",
                    message=str(e),
                    event_type="workflow.failed",
                )
            await _persist_events(task_id)
            await _persist_workflow(task_id)
            raise

    background_tasks.add_task(run_workflow)
    async with workflow_lock:
        _append_workflow_event(
            task_id=task_id,
            agent="workflow_engine",
            message="专利申请流程已启动",
            event_type="workflow.started",
        )
    await _persist_events(task_id)
    return {"task_id": task_id, "status": "started"}


@router.post("/workflows/{task_id}/resume")
async def resume_workflow(task_id: str, background_tasks: BackgroundTasks):
    """恢复工作流执行
    - 中断的流程从当前阶段继续执行
    - 已完成的流程自动进入迭代修正（从撰写阶段重新开始）
    """
    async with workflow_lock:
        context = workflow_engine.get_workflow(task_id)
    if not context:
        raise HTTPException(status_code=404, detail="工作流不存在")

    force_start_from = None
    if context.current_phase == WorkflowState.COMPLETED:
        # 已完成的工作流 → 进入迭代修正模式
        if context.iteration_count >= context.max_iterations:
            raise HTTPException(
                status_code=400,
                detail=f"已达到最大迭代次数（{context.max_iterations}），无法继续修正",
            )
        force_start_from = WorkflowState.PATENT_WRITING
        context.iteration_count += 1

    elif context.current_phase in [WorkflowState.FAILED, WorkflowState.CANCELLED]:
        raise HTTPException(
            status_code=400,
            detail=f"工作流已终止，无法恢复（状态: {context.current_phase.value}）",
        )

    async def phase_callback(phase, result):
        async with workflow_lock:
            _append_workflow_event(
                task_id=task_id,
                agent=phase.value,
                message=f"阶段 {phase.value} 已完成",
                event_type="workflow.phase.completed",
                data={
                    "phase": phase.value,
                    "success": result.success,
                    "duration_seconds": result.duration_seconds,
                    "issues": result.issues,
                },
            )

    async def run_resume():
        try:
            await workflow_engine.resume_workflow(
                context,
                phase_callback=phase_callback,
                force_start_from=force_start_from,
            )
            async with workflow_lock:
                _append_workflow_event(
                    task_id=task_id,
                    agent="workflow_engine",
                    message="专利工作流已恢复并完成",
                    event_type="workflow.completed",
                    data={"state": context.current_phase.value},
                )
            await _persist_events(task_id)
        except Exception as e:
            async with workflow_lock:
                _append_workflow_event(
                    task_id=task_id,
                    agent="workflow_engine",
                    message=str(e),
                    event_type="workflow.failed",
                )
            await _persist_events(task_id)
            raise

    background_tasks.add_task(run_resume)
    async with workflow_lock:
        _append_workflow_event(
            task_id=task_id,
            agent="workflow_engine",
            message=f"工作流已从 {context.current_phase.value} 阶段恢复",
            event_type="workflow.resumed",
            data={"current_phase": context.current_phase.value},
        )
    await _persist_events(task_id)
    return {
        "task_id": task_id,
        "status": "resumed",
        "current_phase": context.current_phase.value,
    }


@router.get("/workflows", response_model=WorkflowListResponse)
async def list_workflows(
    user_id: str = None,
    limit: int = Query(default=50, ge=1, le=100),
    offset: int = Query(default=0, ge=0),
):
    """列出 Hermes/Profile 工作流会话"""
    async with workflow_lock:
        contexts = workflow_engine.list_workflows()

        if user_id:
            contexts = [context for context in contexts if context.user_id == user_id]

        contexts.sort(key=lambda context: context.created_at, reverse=True)
        total = len(contexts)
        items = contexts[offset:offset + limit]
        responses = [_workflow_context_to_response(context) for context in items]

    return WorkflowListResponse(total=total, items=responses)


@router.get("/workflows/{task_id}")
async def get_workflow(task_id: str):
    """获取 Hermes/Profile 工作流状态与阶段输出"""
    async with workflow_lock:
        context = workflow_engine.get_workflow(task_id)
        if not context:
            raise HTTPException(status_code=404, detail="工作流不存在")
        return _workflow_context_to_response(context)


@router.get("/workflows/{task_id}/messages")
async def get_workflow_messages(task_id: str):
    """获取头脑风暴对话历史"""
    async with workflow_lock:
        context = workflow_engine.get_workflow(task_id)
        if not context:
            raise HTTPException(status_code=404, detail="工作流不存在")
        return {"messages": context.message_history, "count": len(context.message_history)}


@router.post("/workflows/{task_id}/cancel")
async def cancel_workflow(task_id: str):
    """取消 Hermes/Profile 工作流"""
    async with workflow_lock:
        if not workflow_engine.cancel_workflow(task_id):
            raise HTTPException(status_code=404, detail="工作流不存在")
        _append_workflow_event(
            task_id=task_id,
            agent="workflow_engine",
            message="工作流已取消",
            event_type="workflow.cancelled",
        )
    await _persist_events(task_id)
    return {"task_id": task_id, "status": "cancelled"}


@router.get("/workflows/{task_id}/stream")
async def stream_workflow_events(task_id: str):
    """SSE 实时事件流 — 工作流各阶段 Agent 的思考/调度/输出事件"""
    context = workflow_engine.get_workflow(task_id)
    if not context:
        raise HTTPException(status_code=404, detail="工作流不存在")

    async def event_generator():
        last_sent = 0
        while True:
            async with workflow_lock:
                events = list(task_events.get(task_id, []))

            # 发送新事件
            for event in events[last_sent:]:
                yield f"event: {event.event_type}\ndata: {event.json()}\n\n"
                last_sent += 1

            # 检查是否结束
            current_phase = context.current_phase.value
            if current_phase in ("completed", "failed", "cancelled"):
                yield f"event: done\ndata: {current_phase}\n\n"
                break

            await asyncio.sleep(1)

    return StreamingResponse(
        event_generator(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "Connection": "keep-alive"},
    )


@router.post("/tasks", response_model=TaskResponse, status_code=status.HTTP_201_CREATED)
async def create_task(request: CreateTaskRequest, background_tasks: BackgroundTasks):
    """创建新的专利申请任务"""
    task_id = str(uuid.uuid4())

    task = PatentTask(
        task_id=task_id,
        user_id=request.user_id,
        tech_description=request.tech_description,
        patent_type_preference=request.patent_type_preference,
        current_state=WorkflowState.INITIAL,
        created_at=datetime.now(),
        updated_at=datetime.now(),
        iteration_count=0,
        max_iterations=3,
    )

    async with workflow_lock:
        tasks_store[task_id] = task
        task_events[task_id] = []

    # 持久化到数据库
    await _persist_task(task_id)
    await _persist_events(task_id)

    logger.info(f"创建新任务: {task_id}")

    # 后台执行工作流
    background_tasks.add_task(_execute_workflow, task_id)

    return task


async def _execute_workflow(task_id: str):
    """在后台执行工作流 — 通过 WorkflowEngine 依次调度各专业 Agent"""
    try:
        async with workflow_lock:
            task = tasks_store[task_id]

        # 1. 创建 WorkflowContext（如果不存在）
        context = workflow_engine.get_workflow(task_id)
        if not context:
            context = workflow_engine.create_workflow(
                task_id=task_id,
                user_id=task.user_id,
                description=task.tech_description,
                patent_type_preference=task.patent_type_preference,
            )

        # 2. 事件回调：将 agent 事件写入 task_events 供 SSE stream 读取
        def _workflow_event_callback(
            agent_name: str, event_type: str, message: str, data: Dict[str, Any]
        ):
            task_events.setdefault(task_id, []).append(
                WorkflowEventResponse(
                    task_id=task_id,
                    timestamp=datetime.now(),
                    agent=agent_name,
                    message=message,
                    event_type=event_type,
                    data=data,
                )
            )

        # 3. 阶段完成回调
        async def phase_callback(phase, result):
            async with workflow_lock:
                _append_workflow_event(
                    task_id=task_id,
                    agent=phase.value,
                    message=f"阶段 {phase.value} 已完成",
                    event_type="workflow.phase.completed",
                    data={
                        "phase": phase.value,
                        "success": result.success,
                        "duration_seconds": result.duration_seconds,
                        "issues": result.issues,
                    },
                )
            await _persist_events(task_id)

        # 4. 执行完整工作流（依次调度各 Agent）
        await workflow_engine.execute_full_workflow(
            context,
            phase_callback=phase_callback,
            event_callback=_workflow_event_callback,
        )

        # 5. 同步状态到 tasks_store
        async with workflow_lock:
            _append_workflow_event(
                task_id=task_id,
                agent="workflow_engine",
                message="专利申请流程已完成",
                event_type="workflow.completed",
                data={"state": context.current_phase.value},
            )
            tasks_store[task_id].current_state = WorkflowState.COMPLETED
        await _persist_task(task_id)
        await _persist_events(task_id)
        await _persist_workflow(task_id)
        logger.info(f"任务 {task_id} 工作流执行完成")

    except Exception as e:
        logger.exception(f"任务 {task_id} 执行失败: {e}")
        async with workflow_lock:
            if task_id in tasks_store:
                tasks_store[task_id].current_state = WorkflowState.FAILED
                tasks_store[task_id].error_message = str(e)
            _append_workflow_event(
                task_id=task_id,
                agent="workflow_engine",
                message=str(e),
                event_type="workflow.failed",
            )
        await _persist_task(task_id)
        await _persist_events(task_id)


@router.get("/tasks", response_model=TaskListResponse)
async def list_tasks(
    user_id: str = None,
    limit: int = Query(default=20, ge=1, le=100),
    offset: int = Query(default=0, ge=0),
):
    """列出所有任务"""
    async with workflow_lock:
        tasks = list(tasks_store.values())

    if user_id:
        tasks = [t for t in tasks if t.user_id == user_id]

    # 按创建时间倒序
    tasks.sort(key=lambda t: t.created_at, reverse=True)

    return {
        "total": len(tasks),
        "tasks": tasks[offset:offset + limit],
    }


@router.get("/tasks/{task_id}", response_model=TaskDetailResponse)
async def get_task(task_id: str):
    """获取任务详情"""
    async with workflow_lock:
        task = tasks_store.get(task_id)
    if not task:
        raise HTTPException(status_code=404, detail="任务不存在")

    # 构建响应 - 转换各阶段文档
    response = TaskDetailResponse(
        task_id=task.task_id,
        user_id=task.user_id,
        current_state=task.current_state,
        created_at=task.created_at,
        updated_at=task.updated_at,
        iteration_count=task.iteration_count,
        error_message=task.error_message,
    )

    # 序列化各阶段输出
    if task.requirement_doc:
        response.requirement_doc = task.requirement_doc.model_dump()

    if task.retrieval_report:
        response.retrieval_report = task.retrieval_report.model_dump()

    if task.draft_doc:
        response.draft_doc = task.draft_doc.model_dump()

    if task.review_report:
        response.review_report = task.review_report.model_dump()

    if task.final_patent:
        response.final_patent = task.final_patent if isinstance(task.final_patent, dict) else task.final_patent.model_dump()

    return response


@router.get("/tasks/{task_id}/events", response_model=List[WorkflowEventResponse])
async def get_task_events(task_id: str):
    """获取任务事件流"""
    async with workflow_lock:
        events = task_events.get(task_id)
    if events is None:
        raise HTTPException(status_code=404, detail="任务不存在")

    return events


@router.get("/tasks/{task_id}/stream")
async def stream_task_events(task_id: str):
    """SSE实时事件流"""
    async with workflow_lock:
        if task_id not in tasks_store:
            raise HTTPException(status_code=404, detail="任务不存在")

    async def event_generator():
        last_sent = 0  # 从头回放所有事件（包含子agent的thinking/tool_call）
        while True:
            async with workflow_lock:
                events = list(task_events.get(task_id, []))
                task = tasks_store.get(task_id)

            # 发送新事件
            for event in events[last_sent:]:
                yield f"event: {event.event_type}\ndata: {event.json()}\n\n"
                last_sent += 1

            # 检查是否结束
            if task and task.current_state in [WorkflowState.COMPLETED, WorkflowState.FAILED]:
                yield f"event: done\ndata: {task.current_state}\n\n"
                break

            await asyncio.sleep(1)

    return StreamingResponse(
        event_generator(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "Connection": "keep-alive"},
    )


@router.post("/tasks/{task_id}/cancel")
async def cancel_task(task_id: str):
    """取消任务"""
    async with workflow_lock:
        task = tasks_store.get(task_id)
        if not task:
            raise HTTPException(status_code=404, detail="任务不存在")

        if task.current_state not in [WorkflowState.COMPLETED, WorkflowState.FAILED]:
            task.current_state = WorkflowState.FAILED
            task.error_message = "用户取消任务"
            logger.info(f"任务 {task_id} 已被用户取消")

    await _persist_task(task_id)
    return {"status": "success", "message": "任务已取消"}


@router.post("/search/patents", response_model=SearchResponse)
async def search_patents(request: SearchPatentRequest):
    """搜索现有技术专利"""
    import time
    start_time = time.time()

    data_source_manager = get_data_source_manager()
    results = await data_source_manager.search_all(request)

    response_results = [
        PriorArtReferenceResponse(
            reference_id=r.reference_id,
            title=r.title,
            publication_date=r.publication_date,
            applicant=r.applicant,
            abstract=r.abstract,
            similarity_score=r.similarity_score,
            source=r.source,
            url=r.url,
        )
        for r in results
    ]

    return SearchResponse(
        total=len(results),
        results=response_results,
        query=request.query,
        search_time=time.time() - start_time,
    )


@router.get("/knowledge/search", response_model=KnowledgeBaseSearchResponse)
async def search_knowledge_base(query: str, top_k: int = 5):
    """搜索本地知识库中的专利"""
    kb = get_knowledge_base()
    patents = kb.search_similar(query, top_k)

    return KnowledgeBaseSearchResponse(
        total=len(patents),
        patents=patents,
        query=query,
    )


@router.get("/system/status", response_model=SystemStatusResponse)
async def get_system_status():
    """获取系统状态"""
    async with workflow_lock:
        active_tasks = sum(
            1 for t in tasks_store.values()
            if t.current_state not in [WorkflowState.COMPLETED, WorkflowState.FAILED]
        )

    kb = get_knowledge_base()

    return SystemStatusResponse(
        status="running",
        active_tasks=active_tasks,
        agents=[
            {"name": "CEO Agent", "description": "全局流程调度", "status": "idle"},
            {"name": "需求分析Agent", "description": "技术需求结构化", "status": "idle"},
            {"name": "检索分析Agent", "description": "专利性评估", "status": "idle"},
            {"name": "专利撰写Agent", "description": "申请文件生成", "status": "idle"},
            {"name": "质量审查Agent", "description": "合规性检查", "status": "idle"},
        ],
        knowledge_base_count=len(kb.list_all_patents()),
        data_sources=["uspto", "epo", "cnipa", "google_patents", "arxiv"],
    )


@router.get("/health")
async def health_check():
    """健康检查"""
    return {"status": "healthy", "timestamp": datetime.now().isoformat()}


# ============ 聊天消息相关 ============
# 存储会话状态
chat_sessions = {}

def generate_brainstorm_response(content: str, phase: str = "initial"):
    """生成头脑风暴阶段的智能回复"""

    responses = {
        "initial": f"""感谢您的描述！为了更准确地评估您的专利申请方案，我想进一步了解几个关键问题：

**1. 技术领域确认**
您的发明具体属于哪个细分技术领域？（例如：人工智能/自然语言处理、半导体/芯片设计、机械/自动化设备等）

**2. 现有技术痛点**
目前行业中针对这个问题的现有解决方案有什么不足？您的发明主要改进了哪些方面？

您可以先回答这两个问题，我们逐步完善信息。""",

        "questioning": f"""非常好！让我们继续完善信息：

**1. 核心创新点**
您认为这个发明最核心的创新点是什么？（可以列出1-3个关键点）

**2. 技术实现细节**
能否简要说明一下技术实现的关键步骤或原理？""",

        "summarizing": f"""太棒了！感谢您的详细说明。根据我们的沟通，我为您整理了专利申请方案摘要：

📋 **专利申请方案摘要**

**技术领域：** 待确认
**核心问题：** 待确认
**创新亮点：**
• 创新点1
• 创新点2
• 创新点3

**技术优势：**
相较于现有技术，您的发明具有以下显著优势：
- 解决了行业长期存在的技术痛点
- 技术方案具备新颖性和创造性
- 具有明确的商业化应用前景

**建议专利类型：** 发明专利

---

请您确认以上信息是否准确？如有需要补充或修改的地方，请随时告诉我。确认无误后，我们可以启动正式的专利申请流程！""",
    }

    return responses.get(phase, responses["initial"])


@router.post("/chat/messages")
async def send_chat_message(request: ChatMessageRequest, phase: str = "initial"):
    """发送聊天消息 - 支持头脑风暴阶段"""
    message_id = str(uuid.uuid4())
    timestamp = datetime.now().isoformat()

    # 如果是已有 Hermes 工作流，直接交给头脑风暴 Agent
    async with workflow_lock:
        workflow_context = workflow_engine.get_workflow(request.task_id) if request.task_id else None
    if workflow_context:
        response = await workflow_engine.add_chat_message(task_id=request.task_id, role="user", content=request.content)
        assistant_content = response.get("content", "消息已记录")
    else:
        # 头脑风暴模式
        assistant_content = generate_brainstorm_response(request.content, phase)

    assistant_response = {
        "id": str(uuid.uuid4()),
        "role": "assistant",
        "content": assistant_content,
        "timestamp": timestamp,
        "phase": phase,
    }

    return {
        "user_message": {
            "id": message_id,
            "role": "user",
            "content": request.content,
            "timestamp": timestamp,
        },
        "assistant_message": assistant_response,
    }


@router.get("/chat/messages")
async def get_chat_messages(session_id: str = "default", task_id: str = None):
    """获取聊天历史"""
    async with workflow_lock:
        context = workflow_engine.get_workflow(task_id) if task_id else None
        if context:
            return {"messages": context.message_history, "count": len(context.message_history)}

    return {
        "messages": [
            {
                "id": "1",
                "role": "assistant",
                "content": "您好！我是专利智脑的智能助手。我将协助您完成专利申请的全过程。请您描述一下您的发明创造。",
                "timestamp": datetime.now().isoformat(),
            }
        ],
        "count": 1,
    }


# ============ Agent管理相关 ============
@router.get("/agents", response_model=AgentListResponse)
async def list_agents() -> AgentListResponse:
    """列出所有 Agent（基于真实 hermes-agent 配置）"""
    now = datetime.now().isoformat()
    agents = []

    # 从 HermesAgentService 获取真实配置
    for cfg in _hermes_service.get_all_configs():
        agents.append({
            "id": cfg.profile_id,
            "name": cfg.name,
            "description": cfg.description,
            "role": _hermes_role_to_ui(cfg.role),
            "system_prompt": cfg.soul_md[:200] + "..." if len(cfg.soul_md) > 200 else cfg.soul_md,
            "model": cfg.model or "default",
            "temperature": cfg.temperature,
            "max_tokens": cfg.max_tokens,
            "working_directory": f"./hermes_agents/{cfg.dir_path.name}",
            "enabled": True,
            "created_at": now,
            "updated_at": now,
            "parent_id": None,
            "child_ids": [],
        })

    return {"agents": agents, "total": len(agents)}


@router.get("/agents/{agent_id}", response_model=AgentDetailResponse)
async def get_agent_detail(agent_id: str) -> AgentDetailResponse:
    """获取 Agent 详情（基于真实 hermes-agent 配置 + registry 工具）"""
    now = datetime.now().isoformat()

    # 优先从 HermesAgentService 获取
    cfg = _hermes_service.get_config(agent_id)
    if cfg:
        # 确保专利工具已注册
        _hermes_service._ensure_patent_tools()

        # 从 hermes registry 获取真实工具信息
        from tools.registry import registry as hermes_registry
        all_patent_tools = {
            name: entry for name, entry in hermes_registry._tools.items()
            if entry.toolset == "patent"
        }

        # 该 Agent 启用的工具
        agent_tools = []
        for tool_name in cfg.enabled_tools:
            entry = all_patent_tools.get(tool_name)
            is_disabled = _override_store.is_tool_disabled(agent_id, tool_name)
            agent_tools.append({
                "id": tool_name,
                "name": tool_name,
                "description": entry.description if entry else f"Hermes 工具: {tool_name}",
                "enabled": not is_disabled,
                "category": _agent_tool_category(tool_name),
                "config": {},
                "is_hermes": True,
                "source_code": None,
                "related_files": [_KNOWN_TOOL_IMPL_FILES.get(tool_name)] if _KNOWN_TOOL_IMPL_FILES.get(tool_name) else [],
            })

        # 用户添加的额外工具
        added_tools = _override_store.get_added_tools(agent_id)
        agent_tools.extend(added_tools)

        # 从文件系统获取真实技能
        agent_skills = []
        for skill_info in cfg.skills:
            skill_name = skill_info.get("name", "")
            is_disabled = _override_store.is_skill_disabled(agent_id, skill_name)
            agent_skills.append({
                "id": skill_name,
                "name": skill_name,
                "description": skill_info.get("description", ""),
                "enabled": not is_disabled,
                "version": "1.0.0",
                "tags": [],
                "source_code": None,
                "source_markdown": None,
                "related_files": [f"hermes_agents/{cfg.dir_path.name}/skills/{skill_info.get('file', '')}"],
            })

        # 用户添加的额外技能
        added_skills = _override_store.get_added_skills(agent_id)
        agent_skills.extend(added_skills)

        # 定时器
        timers = _override_store.get_timers(agent_id)

        # 记忆（从 hermes-agent session 文件读取实际数据）
        memory_cfg = cfg.config.get("memory", {})
        memories = []
        mem_stats = _get_agent_memory_stats(cfg.profile_id, cfg.dir_path.name)

        if memory_cfg.get("short_term", True):
            memories.append({
                "id": "short_term",
                "type": "short_term",
                "name": "短期对话记忆 (Hermes SessionDB)",
                "size": mem_stats["short_term_size"],
                "item_count": mem_stats["short_term_count"],
                "last_updated": mem_stats["last_updated"] or now,
                "content": None,
                "entries": mem_stats["short_term_entries"],
            })
        if memory_cfg.get("long_term", False):
            memories.append({
                "id": "long_term",
                "type": "long_term",
                "name": "长期记忆 (Hermes Persistent Memory)",
                "size": mem_stats["long_term_size"],
                "item_count": mem_stats["long_term_count"],
                "last_updated": mem_stats["last_updated"] or now,
                "content": None,
                "entries": mem_stats["long_term_entries"],
            })
        if memory_cfg.get("knowledge_base", False):
            memories.append({
                "id": "knowledge_base",
                "type": "knowledge_base",
                "name": "知识库 (Hermes Session Search)",
                "size": mem_stats["knowledge_base_size"],
                "item_count": mem_stats["knowledge_base_count"],
                "last_updated": mem_stats["last_updated"] or now,
                "content": None,
                "entries": mem_stats["knowledge_base_entries"],
            })

        config_overrides = _override_store.get_config_overrides(agent_id)
        config_data = {
            "id": cfg.profile_id,
            "name": cfg.name,
            "description": cfg.description,
            "role": _hermes_role_to_ui(cfg.role),
            "system_prompt": cfg.soul_md,
            "model": config_overrides.get("model", cfg.model) or "default",
            "temperature": config_overrides.get("temperature", cfg.temperature),
            "max_tokens": config_overrides.get("max_tokens", cfg.max_tokens),
            "working_directory": f"./hermes_agents/{cfg.dir_path.name}",
            "enabled": True,
            "created_at": now,
            "updated_at": now,
            "parent_id": None,
            "child_ids": [],
        }

        return {
            "config": config_data,
            "tools": agent_tools,
            "skills": agent_skills,
            "timers": timers,
            "memories": memories,
        }

    # Agent 不存在
    raise HTTPException(status_code=404, detail="Agent不存在")


@router.put("/agents/{agent_id}", response_model=AgentUpdateResponse)
async def update_agent_config(agent_id: str, config: dict) -> AgentUpdateResponse:
    """更新Agent配置（持久化到 override store）"""
    cfg = _require_agent_profile(agent_id)

    # 处理嵌套的 tools/skills/timers 批量更新
    if "tools" in config:
        tools_data = config.pop("tools")
        profile_tool_names = set(cfg.enabled_tools)
        for tool in tools_data:
            tool_name = tool.get("name", "")
            if tool_name in profile_tool_names:
                _override_store.toggle_tool(agent_id, tool_name, tool.get("enabled", True))
            else:
                existing_added = _override_store.get_added_tools(agent_id)
                if not any(t.get("id") == tool.get("id") for t in existing_added):
                    _override_store.add_tool(agent_id, tool)

    if "skills" in config:
        skills_data = config.pop("skills")
        cfg_skill_names = {s.get("name", "") for s in cfg.skills}
        for skill in skills_data:
            skill_name = skill.get("name", "")
            if skill_name in cfg_skill_names:
                _override_store.toggle_skill(agent_id, skill_name, skill.get("enabled", True))
            else:
                existing_added = _override_store.get_added_skills(agent_id)
                if not any(s.get("id") == skill.get("id") for s in existing_added):
                    _override_store.add_skill(agent_id, skill)

    if "timers" in config:
        timers_data = config.pop("timers")
        entry = _override_store._ensure_agent(agent_id)
        entry["timers"] = timers_data
        _override_store.save()

    # 剩余字段作为 config overrides 保存
    if config:
        _override_store.update_config(agent_id, config)

    return {
        "agent_id": agent_id,
        "updated_fields": list(config.keys()) + ["tools", "skills", "timers"],
        "message": "Agent配置已持久化更新",
    }


@router.post("/agents/{agent_id}/tools/{tool_id}/toggle", response_model=AgentToggleResponse)
async def toggle_agent_tool(agent_id: str, tool_id: str, enabled: bool) -> AgentToggleResponse:
    """启用/禁用Agent工具（持久化）"""
    _require_agent_profile(agent_id)
    _override_store.toggle_tool(agent_id, tool_id, enabled)
    return {"agent_id": agent_id, "tool_id": tool_id, "enabled": enabled}


@router.post("/agents/{agent_id}/skills/{skill_id}/toggle", response_model=AgentToggleResponse)
async def toggle_agent_skill(agent_id: str, skill_id: str, enabled: bool) -> AgentToggleResponse:
    """启用/禁用Agent技能（持久化）"""
    _require_agent_profile(agent_id)
    _override_store.toggle_skill(agent_id, skill_id, enabled)
    return {"agent_id": agent_id, "skill_id": skill_id, "enabled": enabled}


@router.post("/agents/{agent_id}/timers/{timer_id}/toggle", response_model=AgentToggleResponse)
async def toggle_agent_timer(agent_id: str, timer_id: str, enabled: bool) -> AgentToggleResponse:
    """启用/禁用Agent定时器（持久化）"""
    _require_agent_profile(agent_id)
    success = _override_store.toggle_timer(agent_id, timer_id, enabled)
    if not success:
        raise HTTPException(status_code=404, detail="Agent定时器不存在")
    return {"agent_id": agent_id, "timer_id": timer_id, "enabled": enabled}


@router.post("/agents/{agent_id}/memory/{memory_id}/clear", response_model=AgentToggleResponse)
async def clear_agent_memory(agent_id: str, memory_id: str) -> AgentToggleResponse:
    """清空Agent记忆"""
    _require_agent_profile(agent_id)
    valid_memory_ids = {"short_term", "long_term", "knowledge_base"}
    if memory_id not in valid_memory_ids:
        raise HTTPException(status_code=404, detail="Agent记忆不存在")
    # TODO: 对接真实记忆系统清空操作
    return {"agent_id": agent_id, "memory_id": memory_id, "cleared": True}


# ============ Agent Timer CRUD ============
@router.post("/agents/{agent_id}/timers")
async def create_agent_timer(agent_id: str, timer_data: dict) -> dict:
    """创建新定时器（持久化）"""
    _require_agent_profile(agent_id)
    timer_id = timer_data.get("id") or f"timer_{uuid.uuid4().hex[:8]}"
    timer_entry = {
        "id": timer_id,
        "name": timer_data.get("name", "未命名定时器"),
        "cron_expression": timer_data.get("cron_expression", "0 * * * *"),
        "action": timer_data.get("action", ""),
        "enabled": timer_data.get("enabled", True),
        "last_run": None,
        "next_run": None,
    }
    _override_store.add_timer(agent_id, timer_entry)
    return {"success": True, "timer": timer_entry}


@router.put("/agents/{agent_id}/timers/{timer_id}")
async def update_agent_timer(agent_id: str, timer_id: str, updates: dict) -> dict:
    """更新定时器（持久化）"""
    _require_agent_profile(agent_id)
    success = _override_store.update_timer(agent_id, timer_id, updates)
    if not success:
        raise HTTPException(status_code=404, detail="定时器不存在")
    return {"success": True, "timer_id": timer_id}


@router.delete("/agents/{agent_id}/timers/{timer_id}")
async def delete_agent_timer(agent_id: str, timer_id: str) -> dict:
    """删除定时器（持久化）"""
    _require_agent_profile(agent_id)
    success = _override_store.delete_timer(agent_id, timer_id)
    if not success:
        raise HTTPException(status_code=404, detail="定时器不存在")
    return {"success": True, "timer_id": timer_id}


# ============ Agent Tool CRUD (individual) ============
@router.post("/agents/{agent_id}/tools")
async def create_agent_tool(agent_id: str, tool_data: dict) -> dict:
    """创建新工具（持久化）"""
    _require_agent_profile(agent_id)
    tool_id = tool_data.get("id") or f"tool_{uuid.uuid4().hex[:8]}"
    tool_entry = {
        "id": tool_id,
        "name": tool_data.get("name", ""),
        "description": tool_data.get("description", ""),
        "enabled": tool_data.get("enabled", True),
        "category": tool_data.get("category", "analysis"),
        "config": tool_data.get("config", {}),
    }
    _override_store.add_tool(agent_id, tool_entry)
    return {"success": True, "tool": tool_entry}


@router.put("/agents/{agent_id}/tools/{tool_id}")
async def update_agent_tool(agent_id: str, tool_id: str, tool_data: dict) -> dict:
    """更新工具（持久化）"""
    _require_agent_profile(agent_id)
    success = _override_store.update_tool(agent_id, tool_id, tool_data)
    if not success:
        raise HTTPException(status_code=404, detail="工具不存在")
    return {"success": True, "tool": tool_data}


@router.delete("/agents/{agent_id}/tools/{tool_id}")
async def delete_agent_tool_endpoint(agent_id: str, tool_id: str) -> dict:
    """删除工具（持久化）"""
    _require_agent_profile(agent_id)
    success = _override_store.delete_tool(agent_id, tool_id)
    if not success:
        raise HTTPException(status_code=404, detail="工具不存在")
    return {"success": True, "tool_id": tool_id}


# ============ Agent Skill CRUD (individual) ============
@router.post("/agents/{agent_id}/skills")
async def create_agent_skill(agent_id: str, skill_data: dict) -> dict:
    """创建新技能（持久化）"""
    _require_agent_profile(agent_id)
    skill_id = skill_data.get("id") or f"skill_{uuid.uuid4().hex[:8]}"
    skill_entry = {
        "id": skill_id,
        "name": skill_data.get("name", ""),
        "description": skill_data.get("description", ""),
        "enabled": skill_data.get("enabled", True),
        "version": skill_data.get("version", "1.0.0"),
        "tags": skill_data.get("tags", []),
    }
    _override_store.add_skill(agent_id, skill_entry)
    return {"success": True, "skill": skill_entry}


@router.delete("/agents/{agent_id}/skills/{skill_id}")
async def delete_agent_skill_endpoint(agent_id: str, skill_id: str) -> dict:
    """删除技能（持久化）"""
    _require_agent_profile(agent_id)
    success = _override_store.delete_skill(agent_id, skill_id)
    if not success:
        raise HTTPException(status_code=404, detail="技能不存在")
    return {"success": True, "skill_id": skill_id}


# ============ Hermes Agent 对话 (基于真实 hermes-agent) ============
@router.post("/agents/{agent_id}/chat")
async def hermes_agent_chat(agent_id: str, body: dict) -> dict:
    """使用真实 hermes-agent AIAgent 进行对话"""
    content = body.get("content", "").strip()
    session_id = body.get("session_id")
    user_id = body.get("user_id", "default_user")

    if not content:
        raise HTTPException(status_code=400, detail="对话内容不能为空")

    # 使用 HermesAgentService 运行对话
    try:
        result = await _hermes_service.run_conversation(
            profile_id=agent_id,
            user_input=content,
            session_id=session_id,
            user_id=user_id,
        )
        # hermes-agent 返回字典，提取 final_response
        if isinstance(result, dict):
            response_text = result.get("final_response", "") or ""
        else:
            response_text = str(result) if result else ""

        return {
            "success": True,
            "agent_id": agent_id,
            "response": response_text,
            "session_id": session_id,
            "metadata": {
                "model": result.get("model") if isinstance(result, dict) else None,
                "input_tokens": result.get("input_tokens") if isinstance(result, dict) else None,
                "output_tokens": result.get("output_tokens") if isinstance(result, dict) else None,
            },
        }
    except ValueError as e:
        raise HTTPException(status_code=404, detail=str(e))
    except Exception as e:
        logger.error(f"Hermes agent chat failed: {e}")
        raise HTTPException(status_code=500, detail=f"Agent 对话失败: {str(e)[:200]}")


@router.post("/agents/{agent_id}/chat/stream")
async def hermes_agent_chat_stream(agent_id: str, body: dict):
    """使用真实 hermes-agent AIAgent 进行流式对话（SSE）"""
    import json as _json

    content = body.get("content", "").strip()
    session_id = body.get("session_id")
    user_id = body.get("user_id", "default_user")

    if not content:
        raise HTTPException(status_code=400, detail="对话内容不能为空")

    async def event_generator():
        try:
            async for event in _hermes_service.run_conversation_stream(
                profile_id=agent_id,
                user_input=content,
                session_id=session_id,
                user_id=user_id,
            ):
                event_type = event.get("type", "status")
                event_data = event.get("data", {})
                yield f"event: {event_type}\ndata: {_json.dumps(event_data, ensure_ascii=False)}\n\n"
        except ValueError as e:
            yield f"event: error\ndata: {_json.dumps({'error': str(e)}, ensure_ascii=False)}\n\n"
        except Exception as e:
            yield f"event: error\ndata: {_json.dumps({'error': str(e)[:200]}, ensure_ascii=False)}\n\n"

    return StreamingResponse(
        event_generator(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "Connection": "keep-alive", "X-Accel-Buffering": "no"},
    )


@router.get("/agents/{agent_id}/hermes-info")
async def get_hermes_agent_info(agent_id: str) -> dict:
    """获取 Hermes Agent 的运行时信息（基于真实 hermes-agent）"""
    config = _hermes_service.get_config(agent_id)
    if not config:
        raise HTTPException(status_code=404, detail=f"Hermes Agent 未找到: {agent_id}")

    return {
        "profile_id": config.profile_id,
        "name": config.name,
        "description": config.description,
        "role": config.role,
        "model": config.model,
        "temperature": config.temperature,
        "max_tokens": config.max_tokens,
        "enabled_toolsets": config.enabled_toolsets,
        "enabled_tools": config.enabled_tools,
        "skills": config.skills,
        "hermes_agent_version": "0.15.1",
        "capabilities": {
            "tool_calling": True,
            "memory_persistent": True,
            "skills_self_improving": True,
            "cron_scheduling": True,
            "context_compression": True,
            "session_search": True,
        },
    }


# ============ 组织架构相关 ============
@router.get("/organization/tree", response_model=OrgNodeResponse)
async def get_organization_tree() -> OrgNodeResponse:
    """获取组织架构树"""
    if organization_tree_store is not None:
        return organization_tree_store
    return _profile_organization_tree()


@router.put("/organization/tree", response_model=OrganizationUpdateResponse)
async def update_organization_tree(tree: OrgNodeResponse) -> OrganizationUpdateResponse:
    """更新组织架构树"""
    global organization_tree_store
    _validate_org_tree(tree)
    organization_tree_store = tree
    await _persist_org_tree()
    return {"status": "success", "message": "组织架构已更新", "tree": tree}


# ============ 对话会话相关 ============

def _get_brainstorm_agent():
    """获取 CEO Agent 实例（使用真实 hermes-agent AIAgent）"""
    return _hermes_service.create_agent_instance(
        profile_id="patent.ceo.v1",
        session_id=f"conversation_{uuid.uuid4().hex[:8]}",
    )


@router.post("/conversations", response_model=ConversationDetail, status_code=status.HTTP_201_CREATED)
async def create_conversation(request: CreateConversationRequest):
    """创建新的头脑风暴对话会话"""
    conv_id = str(uuid.uuid4())
    now = datetime.now().isoformat()
    greeting_id = str(uuid.uuid4())
    conversation = {
        "id": conv_id,
        "title": request.title or "新的对话",
        "messages": [
            {
                "id": greeting_id,
                "role": "assistant",
                "content": "你好！我是专利智脑的创意助手。请告诉我你的技术构思，我们可以一起探讨它的专利价值。",
                "timestamp": now,
                "type": "text",
                "metadata": None,
            }
        ],
        "created_at": now,
        "updated_at": now,
        "status": "brainstorming",
        "linked_workflow_id": None,
    }
    async with conversations_lock:
        conversations_store[conv_id] = conversation
    await _persist_conversation(conv_id)
    return conversation


@router.get("/conversations", response_model=ConversationListResponse)
async def list_conversations(
    limit: int = Query(default=50, ge=1, le=100),
    offset: int = Query(default=0, ge=0),
):
    """列出所有对话会话"""
    async with conversations_lock:
        items = []
        for conv in conversations_store.values():
            items.append({
                "id": conv["id"],
                "title": conv["title"],
                "created_at": conv["created_at"],
                "updated_at": conv["updated_at"],
                "message_count": len(conv["messages"]),
                "status": conv["status"],
                "linked_workflow_id": conv.get("linked_workflow_id"),
            })
        items.sort(key=lambda x: x["updated_at"], reverse=True)
        total = len(items)
        page_items = items[offset:offset + limit]
    return ConversationListResponse(items=page_items, total=total, page=(offset // limit) + 1, page_size=limit)


@router.get("/conversations/{conv_id}", response_model=ConversationDetail)
async def get_conversation(conv_id: str):
    """获取对话会话详情"""
    async with conversations_lock:
        conv = conversations_store.get(conv_id)
    if not conv:
        raise HTTPException(status_code=404, detail="对话不存在")
    return conv


@router.post("/conversations/{conv_id}/chat", response_model=ConversationChatResponse)
async def chat_in_conversation(conv_id: str, request: ConversationChatRequest):
    """在对话中发送消息（头脑风暴模式，不会创建专利工作流）"""
    content = request.content.strip()
    if not content:
        raise HTTPException(status_code=400, detail="消息内容不能为空")

    async with conversations_lock:
        conv = conversations_store.get(conv_id)
    if not conv:
        raise HTTPException(status_code=404, detail="对话不存在")
    if conv.get("linked_workflow_id"):
        raise HTTPException(status_code=400, detail="该对话已关联工作流，请使用工作流聊天接口")

    # 添加用户消息
    now = datetime.now().isoformat()
    user_msg = {
        "id": str(uuid.uuid4()),
        "role": "user",
        "content": content,
        "timestamp": now,
        "type": "text",
        "metadata": None,
    }

    try:
        agent = _get_brainstorm_agent()
        history_text = "\n".join([
            f"{m['role'].upper()}: {m['content']}"
            for m in conv["messages"][-10:]
        ])

        prompt = f"""对话历史:
{history_text}

用户: {content}

要求:
- 言简意赅回复，3-5句话
- 需要分析时调用工具（ipc_classifier/risk_analyzer/task_planner/patent_search/tech_feature_extractor），不要自己编造
- 提1-2个关键追问推进讨论
- 收集够信息后加标记 [CREATE_PATENT_RECOMMENDATION]
"""
        # 使用真实 hermes-agent AIAgent（run_conversation 是同步方法）
        response = await asyncio.to_thread(agent.run_conversation, prompt)
        # hermes-agent 返回字典或字符串
        if isinstance(response, dict):
            response_text = response.get("final_response", "") or str(response)
        else:
            response_text = str(response) if response else ""

        # hermes-agent 的工具调用记录在 agent 内部，无需手动收集
        tool_calls_data = []
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"AI 响应失败: {str(e)}")

    # 检查 AI 是否建议创建专利
    has_recommendation = "[CREATE_PATENT_RECOMMENDATION]" in response_text
    clean_text = response_text.replace("[CREATE_PATENT_RECOMMENDATION]", "").strip()

    assistant_msg = {
        "id": str(uuid.uuid4()),
        "role": "assistant",
        "content": clean_text,
        "timestamp": datetime.now().isoformat(),
        "type": "text",
        "metadata": {"recommend_create_patent": has_recommendation} if has_recommendation else None,
        "tool_calls": tool_calls_data if tool_calls_data else None,
    }

    async with conversations_lock:
        conv = conversations_store.get(conv_id)
        if conv:
            conv["messages"].append(user_msg)
            conv["messages"].append(assistant_msg)
            conv["updated_at"] = datetime.now().isoformat()
            # 自动生成标题（从第一条用户消息）
            user_msgs = [m for m in conv["messages"] if m["role"] == "user"]
            if conv["title"] == "新的对话" and len(user_msgs) == 1:
                title = content[:50]
                conv["title"] = title + ("..." if len(content) > 50 else "")

    await _persist_conversation(conv_id)
    return ConversationChatResponse(
        message=assistant_msg,
        has_recommendation=has_recommendation,
        conversation_id=conv_id,
    )


@router.post("/conversations/{conv_id}/chat/stream")
async def chat_in_conversation_stream(conv_id: str, request: ConversationChatRequest):
    """在对话中发送消息 — SSE 流式输出（展示工具调用和技能使用过程）"""
    content = request.content.strip()
    if not content:
        raise HTTPException(status_code=400, detail="消息内容不能为空")

    async with conversations_lock:
        conv = conversations_store.get(conv_id)
    if not conv:
        raise HTTPException(status_code=404, detail="对话不存在")
    if conv.get("linked_workflow_id"):
        raise HTTPException(status_code=400, detail="该对话已关联工作流，请使用工作流聊天接口")

    # 添加用户消息
    now = datetime.now().isoformat()
    user_msg = {
        "id": str(uuid.uuid4()),
        "role": "user",
        "content": content,
        "timestamp": now,
        "type": "text",
        "metadata": None,
    }

    async with conversations_lock:
        conv = conversations_store.get(conv_id)
        if conv:
            conv["messages"].append(user_msg)
            conv["updated_at"] = now

    async def event_generator():
        import json as _json

        agent = _get_brainstorm_agent()
        history_text = "\n".join([
            f"{m['role'].upper()}: {m['content']}"
            for m in conv["messages"][-10:]
        ])

        prompt = f"""对话历史:
{history_text}

用户: {content}

要求:
- 言简意赅回复，3-5句话
- 需要分析时调用工具（ipc_classifier/risk_analyzer/task_planner/patent_search/tech_feature_extractor），不要自己编造
- 提1-2个关键追问推进讨论
- 收集够信息后加标记 [CREATE_PATENT_RECOMMENDATION]
"""
        tool_calls_data = []
        skill_uses_data = []
        final_content = ""

        try:
            # 使用真实 hermes-agent 流式对话
            async for event in _hermes_service.run_conversation_stream(
                profile_id="patent.ceo.v1",
                user_input=prompt,
                session_id=f"conv_{conv_id}",
            ):
                event_type = event["type"]
                event_data = event["data"]

                if event_type == "thinking":
                    yield f"event: thinking\ndata: {_json.dumps(event_data, ensure_ascii=False)}\n\n"

                elif event_type == "tool_call_start":
                    yield f"event: tool_call_start\ndata: {_json.dumps(event_data, ensure_ascii=False)}\n\n"

                elif event_type == "tool_call_end":
                    tool_calls_data.append(event_data)
                    yield f"event: tool_call_end\ndata: {_json.dumps(event_data, ensure_ascii=False)}\n\n"

                elif event_type == "content":
                    final_content = event_data.get("content", "")
                    has_recommendation = "[CREATE_PATENT_RECOMMENDATION]" in final_content
                    clean_content = final_content.replace("[CREATE_PATENT_RECOMMENDATION]", "").strip()
                    yield f"event: content\ndata: {_json.dumps({'content': clean_content, 'has_recommendation': has_recommendation}, ensure_ascii=False)}\n\n"

                elif event_type == "done":
                    # 构建 assistant 消息并持久化
                    clean_content = final_content.replace("[CREATE_PATENT_RECOMMENDATION]", "").strip()
                    has_recommendation = "[CREATE_PATENT_RECOMMENDATION]" in final_content

                    assistant_msg = {
                        "id": str(uuid.uuid4()),
                        "role": "assistant",
                        "content": clean_content,
                        "timestamp": datetime.now().isoformat(),
                        "type": "text",
                        "metadata": {"recommend_create_patent": has_recommendation} if has_recommendation else None,
                        "tool_calls": tool_calls_data if tool_calls_data else None,
                        "skill_uses": skill_uses_data if skill_uses_data else None,
                    }

                    async with conversations_lock:
                        c = conversations_store.get(conv_id)
                        if c:
                            c["messages"].append(assistant_msg)
                            c["updated_at"] = datetime.now().isoformat()
                            user_msgs = [m for m in c["messages"] if m["role"] == "user"]
                            if c["title"] == "新的对话" and len(user_msgs) == 1:
                                title = content[:50]
                                c["title"] = title + ("..." if len(content) > 50 else "")

                    await _persist_conversation(conv_id)

                    yield f"event: done\ndata: {_json.dumps({'message': assistant_msg, 'has_recommendation': has_recommendation, 'conversation_id': conv_id}, ensure_ascii=False)}\n\n"

        except Exception as e:
            yield f"event: error\ndata: {_json.dumps({'error': str(e)}, ensure_ascii=False)}\n\n"

    return StreamingResponse(
        event_generator(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "Connection": "keep-alive", "X-Accel-Buffering": "no"},
    )


@router.post("/conversations/{conv_id}/create-workflow", response_model=dict)
async def create_workflow_from_conversation(conv_id: str, request: CreateWorkflowFromConversationRequest, background_tasks: BackgroundTasks):
    """从对话内容创建专利申请工作流（自动启动）"""
    async with conversations_lock:
        conv = conversations_store.get(conv_id)
    if not conv:
        raise HTTPException(status_code=404, detail="对话不存在")
    if conv.get("linked_workflow_id"):
        raise HTTPException(status_code=400, detail="该对话已关联工作流")

    # 从对话中提取技术描述（取所有用户消息合并）
    user_msgs = [m["content"] for m in conv["messages"] if m["role"] == "user"]
    if not user_msgs:
        raise HTTPException(status_code=400, detail="对话中没有用户消息，无法创建工作流")
    tech_description = " ".join(user_msgs)

    # 创建工作流
    async with workflow_lock:
        context = workflow_engine.create_workflow(
            task_id=str(uuid.uuid4()),
            user_id=request.user_id,
            description=tech_description,
        )
        task_id = context.task_id
        task_events[task_id] = []
        _append_workflow_event(
            task_id=task_id,
            agent="workflow_engine",
            message="已从对话内容创建专利工作流",
            event_type="workflow.created",
            data={"state": context.current_phase.value},
        )

    # 关联对话到工作流
    async with conversations_lock:
        conv = conversations_store.get(conv_id)
        if conv:
            conv["linked_workflow_id"] = task_id
            conv["status"] = "workflow_linked"
            conv["updated_at"] = datetime.now().isoformat()

    await _persist_events(task_id)
    await _persist_conversation(conv_id)

    # 自动启动工作流（后台执行）
    async def auto_start_workflow():
        """自动启动工作流的后台任务"""
        def _workflow_event_callback(agent_name: str, event_type: str, message: str, data: Dict[str, Any]):
            """直接将agent事件写入task_events"""
            task_events.setdefault(task_id, []).append(
                WorkflowEventResponse(
                    task_id=task_id,
                    timestamp=datetime.now(),
                    agent=agent_name,
                    message=message,
                    event_type=event_type,
                    data=data,
                )
            )

        async def phase_callback(phase, result):
            async with workflow_lock:
                _append_workflow_event(
                    task_id=task_id,
                    agent=phase.value,
                    message=f"阶段 {phase.value} 已完成",
                    event_type="workflow.phase.completed",
                    data={
                        "phase": phase.value,
                        "success": result.success,
                        "duration_seconds": result.duration_seconds,
                        "issues": result.issues,
                    },
                )
            await _persist_events(task_id)

        try:
            await workflow_engine.execute_full_workflow(
                context,
                phase_callback=phase_callback,
                event_callback=_workflow_event_callback,
            )
            async with workflow_lock:
                _append_workflow_event(
                    task_id=task_id,
                    agent="workflow_engine",
                    message="专利申请流程已完成",
                    event_type="workflow.completed",
                    data={"state": context.current_phase.value},
                )
            await _persist_events(task_id)
            await _persist_workflow(task_id)
        except Exception as e:
            logger.error(f"工作流自动执行失败: {e}")
            async with workflow_lock:
                _append_workflow_event(
                    task_id=task_id,
                    agent="workflow_engine",
                    message=str(e),
                    event_type="workflow.failed",
                )
            await _persist_events(task_id)
            await _persist_workflow(task_id)
            # 将失败信息写入对话，让 CEO 与用户沟通
            async with conversations_lock:
                c = conversations_store.get(conv_id)
                if c:
                    c["messages"].append({
                        "id": str(uuid.uuid4()),
                        "role": "assistant",
                        "content": f"⚠️ 专利申请流程遇到问题：{str(e)}\n\n请补充更多技术细节后重试，或联系我协助解决。",
                        "timestamp": datetime.now().isoformat(),
                        "type": "text",
                        "metadata": {"type": "workflow_error", "error": str(e)},
                    })
            await _persist_conversation(conv_id)

    background_tasks.add_task(auto_start_workflow)

    return {"task_id": task_id, "status": "started", "conversation_id": conv_id}


@router.delete("/conversations/{conv_id}", status_code=status.HTTP_204_NO_CONTENT)
async def delete_conversation(conv_id: str):
    """删除对话会话"""
    async with conversations_lock:
        if conv_id not in conversations_store:
            raise HTTPException(status_code=404, detail="对话不存在")
        del conversations_store[conv_id]
    store = _get_persist_store()
    await store.delete("conversation", conv_id)


@router.get("/conversations/{conv_id}/workflow-status", response_model=dict)
async def get_conversation_workflow_status(conv_id: str):
    """获取对话关联的工作流状态"""
    async with conversations_lock:
        conv = conversations_store.get(conv_id)
    if not conv:
        raise HTTPException(status_code=404, detail="对话不存在")

    linked_id = conv.get("linked_workflow_id")
    if not linked_id:
        return {"linked": False, "workflow_id": None, "status": None}

    async with workflow_lock:
        context = workflow_engine.get_workflow(linked_id)
    if not context:
        return {"linked": True, "workflow_id": linked_id, "status": "not_found"}

    return {
        "linked": True,
        "workflow_id": linked_id,
        "status": context.current_phase.value,
        "message_count": len(context.message_history) if hasattr(context, "message_history") else 0,
    }


# ============ 统计数据相关 ============
@router.get("/stats/dashboard")
async def get_dashboard_stats():
    """获取仪表盘统计数据"""
    async with workflow_lock:
        tasks = list(tasks_store.values())

    return {
        "total_tasks": len(tasks),
        "completed_tasks": sum(1 for t in tasks if t.current_state == WorkflowState.COMPLETED),
        "in_progress_tasks": sum(
            1 for t in tasks
            if t.current_state not in [WorkflowState.COMPLETED, WorkflowState.FAILED]
        ),
        "failed_tasks": sum(1 for t in tasks if t.current_state == WorkflowState.FAILED),
        "active_agents": 5,
        "avg_completion_time": "2.5 hours",
        "success_rate": 94.5,
    }


# ============ Agent 文件浏览与相关文件查看 ============
import os as _os
import stat as _stat
import ast as _ast
import re as _re


def _extract_tool_structure(source_code: str | None, tool_name: str) -> dict:
    """从工具源代码中提取结构元数据"""
    if not source_code:
        return {
            "class_name": None,
            "description": None,
            "parameters": [],
            "methods": [],
            "file_path": _KNOWN_TOOL_IMPL_FILES.get(tool_name),
            "template": _get_tool_template(),
        }

    try:
        tree = _ast.parse(source_code)
    except SyntaxError:
        return {"class_name": None, "parameters": [], "methods": [], "template": _get_tool_template()}

    class_name = None
    description = None
    parameters = []
    methods = []

    for node in _ast.walk(tree):
        if isinstance(node, _ast.ClassDef):
            # 找到 HermesTool 子类
            for base in node.bases:
                base_name = ""
                if isinstance(base, _ast.Name):
                    base_name = base.id
                elif isinstance(base, _ast.Attribute):
                    base_name = base.attr
                if "Tool" in base_name:
                    class_name = node.name
                    # 提取 docstring
                    if (node.body and isinstance(node.body[0], _ast.Expr)
                            and isinstance(node.body[0].value, _ast.Constant)):
                        description = node.body[0].value.value

                    # 提取方法
                    for item in node.body:
                        if isinstance(item, (_ast.FunctionDef, _ast.AsyncFunctionDef)):
                            args_list = [
                                a.arg for a in item.args.args if a.arg != "self"
                            ]
                            methods.append({
                                "name": item.name,
                                "args": args_list,
                                "is_async": isinstance(item, _ast.AsyncFunctionDef),
                            })
                    break

    # 提取 HermesToolParameter 定义
    param_pattern = _re.compile(
        r'"(\w+)":\s*HermesToolParameter\(\s*'
        r'type="(\w+)",\s*description="([^"]*)"',
        _re.DOTALL,
    )
    for match in param_pattern.finditer(source_code):
        parameters.append({
            "name": match.group(1),
            "type": match.group(2),
            "description": match.group(3),
        })

    return {
        "class_name": class_name,
        "description": description,
        "parameters": parameters,
        "methods": methods,
        "file_path": _KNOWN_TOOL_IMPL_FILES.get(tool_name),
        "template": _get_tool_template(),
    }


def _extract_skill_structure(profile_skill, skill_meta: dict) -> dict:
    """从技能数据中提取结构元数据"""
    return {
        "name": skill_meta.get("name", ""),
        "description": skill_meta.get("description", ""),
        "proficiency": profile_skill.proficiency if profile_skill else 0.8,
        "keywords": profile_skill.keywords if profile_skill else [],
        "version": skill_meta.get("version", "1.0.0"),
        "injection_method": "system_prompt",
        "injection_description": "技能信息被注入到 Agent 的 system_prompt 中，引导 LLM 在回复时运用该技能",
        "template": _get_skill_template(),
    }


def _get_tool_template() -> str:
    """返回创建新工具的模板代码"""
    return '''"""
{工具描述}
"""
from typing import Any, Dict

from ..base import HermesTool, HermesToolDefinition, HermesToolParameter
from src.core.logging import get_logger
from src.core.llm_client import get_llm_service, LLMMessage

logger = get_logger(__name__)


class {ClassName}Tool(HermesTool):
    """{工具描述}"""
    name = "{tool_name}"
    description = "{工具功能描述}"

    def _build_definition(self) -> HermesToolDefinition:
        return HermesToolDefinition(
            name=self.name,
            description=self.description,
            parameters={
                "param1": HermesToolParameter(
                    type="string",
                    description="参数1描述",
                    required=True,
                ),
            },
        )

    async def execute(self, param1: str, **kwargs) -> Dict[str, Any]:
        """执行工具逻辑"""
        logger.info("Executing tool", tool=self.name)
        # 实现工具逻辑...
        return {"result": "...", "tool": self.name}
'''


def _get_skill_template() -> str:
    """返回创建新技能的结构说明"""
    return '''{
  "name": "技能名称",
  "description": "技能详细描述 — 说明该技能能做什么",
  "proficiency": 0.85,  // 熟练度 0.0-1.0
  "keywords": ["关键词1", "关键词2"]  // 用于匹配任务的关键词
}

// 技能工作原理：
// 1. 技能定义被添加到 Agent Profile 的 skills 目录
// 2. Agent 创建时，技能信息被注入到 system_prompt
// 3. LLM 根据技能描述在回复中运用相应能力
// 4. Agent 可以在回复中用 <skill_use> 标记声明使用了哪些技能
'''

_KNOWN_TOOL_IMPL_FILES: dict[str, str] = {
    "task_planner": "backend/src/agents/hermes/tools/task_planner.py",
    "quality_assessor": "backend/src/agents/hermes/tools/quality_assessor.py",
    "report_generator": "backend/src/agents/hermes/tools/report_generator.py",
    "risk_analyzer": "backend/src/agents/hermes/tools/risk_analyzer.py",
    "ipc_classifier": "backend/src/agents/hermes/tools/ipc_classifier.py",
    "tech_feature_extractor": "backend/src/agents/hermes/tools/tech_feature_extractor.py",
    "scenario_miner": "backend/src/agents/hermes/tools/scenario_miner.py",
    "patent_search": "backend/src/agents/hermes/tools/patent_search.py",
    "similarity_analyzer": "backend/src/agents/hermes/tools/similarity_analyzer.py",
    "patentability_scorer": "backend/src/agents/hermes/tools/patentability_scorer.py",
    "claim_drafter": "backend/src/agents/hermes/tools/claim_drafter.py",
    "description_writer": "backend/src/agents/hermes/tools/description_writer.py",
    "terminology_normalizer": "backend/src/agents/hermes/tools/terminology_normalizer.py",
    "support_checker": "backend/src/agents/hermes/tools/support_checker.py",
    "compliance_checker": "backend/src/agents/hermes/tools/compliance_checker.py",
    "claim_quality_analyzer": "backend/src/agents/hermes/tools/claim_quality_analyzer.py",
    "support_verifier": "backend/src/agents/hermes/tools/support_verifier.py",
    "oa_predictor": "backend/src/agents/hermes/tools/oa_predictor.py",
    "creative_thinking": "backend/src/agents/hermes/tools/creative_thinking.py",
    "patent_strategy_guide": "backend/src/agents/hermes/tools/patent_strategy_guide.py",
    "agent_selector": "backend/src/agents/hermes/tools/agent_selector.py",
    "dispatch_specialist": "backend/src/agents/hermes/tools/dispatch_specialist.py",
    "prior_art_comparator": "backend/src/agents/hermes/tools/prior_art_comparator.py",
}


@router.get("/agents/{agent_id}/related-files")
async def get_agent_related_files(
    agent_id: str,
    tool_id: str = None,
    skill_id: str = None,
):
    """获取 Agent 特定工具或技能的相关文件列表及内容"""
    if (tool_id is None and skill_id is None) or (tool_id is not None and skill_id is not None):
        raise HTTPException(
            status_code=400,
            detail="必须且只能提供 tool_id 或 skill_id 参数之一",
        )

    cfg = _require_agent_profile(agent_id)

    _project_root = _os.path.dirname(
        _os.path.dirname(_os.path.dirname(_os.path.dirname(_os.path.abspath(__file__))))
    )

    def _read_file(rel_path: str) -> str | None:
        """安全读取项目内文件"""
        abs_path = _os.path.normpath(_os.path.join(_project_root, rel_path))
        if not abs_path.startswith(_project_root):
            return None
        if not _os.path.isfile(abs_path):
            return None
        try:
            with open(abs_path, "r", encoding="utf-8") as f:
                return f.read()
        except Exception:
            return None

    if tool_id is not None:
        if tool_id not in cfg.enabled_tools:
            added_tools = _override_store.get_added_tools(agent_id)
            if not any(t.get("id") == tool_id or t.get("name") == tool_id for t in added_tools):
                raise HTTPException(status_code=404, detail=f"工具未找到: {tool_id}")

        rel_path = _KNOWN_TOOL_IMPL_FILES.get(tool_id)
        if not rel_path:
            raise HTTPException(status_code=404, detail=f"工具实现文件未配置: {tool_id}")
        source_code = _read_file(rel_path)
        if source_code is None:
            raise HTTPException(status_code=404, detail=f"工具实现文件不存在: {rel_path}")

        # 提取工具结构元数据
        structure_info = _extract_tool_structure(source_code, tool_id)

        return {
            "type": "tool",
            "name": tool_id,
            "source_code": source_code,
            "source_markdown": None,
            "structure": structure_info,
            "files": [{"path": rel_path, "content": source_code}],
        }

    else:
        # 从 hermes config 获取技能信息
        skill_meta = next((s for s in cfg.skills if s.get("name") == skill_id), None)
        if skill_meta is None:
            added_skills = _override_store.get_added_skills(agent_id)
            skill_meta = next((s for s in added_skills if s.get("id") == skill_id), None)
            if skill_meta is None:
                raise HTTPException(status_code=404, detail=f"技能未找到: {skill_id}")

        # 结构元数据
        structure_info = _extract_skill_structure(None, skill_meta)

        # 读取实际的 .md 文件内容
        skill_file = skill_meta.get("file", "")
        if not skill_file:
            raise HTTPException(status_code=404, detail=f"技能文件未配置: {skill_id}")
        skill_path = cfg.dir_path / "skills" / skill_file
        if not skill_path.exists():
            raise HTTPException(status_code=404, detail=f"技能文件不存在: {skill_path}")
        source_markdown = skill_path.read_text(encoding="utf-8")

        return {
            "type": "skill",
            "name": skill_meta.get("name", skill_id),
            "source_code": None,
            "source_markdown": source_markdown,
            "structure": structure_info,
            "files": [
                {
                    "path": f"skills/{skill_meta.get('name', skill_id)}.md",
                    "content": source_markdown,
                },
            ],
        }


# ============ Agent 工具/技能生成与热插拔 ============
import textwrap as _textwrap


@router.post("/agents/{agent_id}/tools/chat-generate")
async def chat_generate_tool(agent_id: str, body: dict):
    """通过 LLM 对话生成标准 HermesTool 子类代码"""
    _require_agent_profile(agent_id)
    name = body.get("name", "").strip()
    description = body.get("description", "").strip()
    parameters = body.get("parameters") or {}

    if not name or not description:
        raise HTTPException(status_code=400, detail="工具名称和描述不能为空")

    # 使用 LLM 生成工具代码
    from src.core.llm_client import get_llm_service, LLMMessage as _LLMMsg

    llm = get_llm_service()

    # 构建 param 描述
    param_desc = "\n".join(f"  - {k}: {v}" for k, v in parameters.items()) if parameters else "  （无额外参数）"

    prompt = f"""你是一个 Python 代码生成专家。请生成一个符合 Hermes Agent 框架的工具类。

要求：
- 工具名称: {name}
- 功能描述: {description}
- 输入参数:
{param_desc}

必须严格遵循以下模板结构（不要添加任何额外的导入或代码）：

```python
\"\"\"
{{工具描述}}
\"\"\"
from typing import Any, Dict

from ..base import HermesTool, HermesToolDefinition, HermesToolParameter
from src.core.logging import get_logger
from src.core.llm_client import get_llm_service, LLMMessage

logger = get_logger(__name__)


class {{ClassName}}Tool(HermesTool):
    \"\"\"{{工具描述}}\"\"\"
    name = "{{tool_name}}"
    description = "{{功能描述}}"

    def _build_definition(self) -> HermesToolDefinition:
        return HermesToolDefinition(
            name=self.name,
            description=self.description,
            parameters={{
                # 每个参数都是 HermesToolParameter
            }},
        )

    async def execute(self, **kwargs) -> Dict[str, Any]:
        \"\"\"执行工具逻辑\"\"\"
        logger.info("Executing tool", tool=self.name)
        # 实现逻辑...
        return {{"result": "...", "tool": self.name}}
```

请生成完整的工具代码，只输出 Python 代码，不要输出其他内容。确保：
1. 类名使用 CamelCase（如 PatentSearch → PatentSearchTool）
2. 参数使用 HermesToolParameter 定义（type, description, required）
3. execute 方法实现真正的功能逻辑（可调用 LLM）
4. 包含 register 函数
"""

    try:
        response = await llm.chat_completion(
            messages=[_LLMMsg(role="user", content=prompt)],
            temperature=0.3,
        )
        generated_code = response.content or ""

        # 清理可能的 markdown 包裹
        if "```python" in generated_code:
            generated_code = generated_code.split("```python", 1)[1]
            if "```" in generated_code:
                generated_code = generated_code.rsplit("```", 1)[0]
        elif "```" in generated_code:
            parts = generated_code.split("```")
            if len(parts) >= 3:
                generated_code = parts[1]

        generated_code = generated_code.strip()

        return {
            "success": True,
            "name": name,
            "code": generated_code,
            "message": "工具代码已通过 LLM 生成，请验证后注册",
        }
    except Exception as e:
        # 降级：生成模板代码
        class_name = "".join(word.capitalize() for word in name.split("_"))
        param_defs = ""
        param_args = ""
        for k, v in parameters.items():
            param_defs += f'''                "{k}": HermesToolParameter(\n                    type="string",\n                    description="{v or k}",\n                    required=True,\n                ),\n'''
            param_args += f"    {k}: str, "

        code = f'''"""
{description}
"""
from typing import Any, Dict

from ..base import HermesTool, HermesToolDefinition, HermesToolParameter
from src.core.logging import get_logger
from src.core.llm_client import get_llm_service, LLMMessage

logger = get_logger(__name__)


class {class_name}Tool(HermesTool):
    """{description}"""
    name = "{name}"
    description = "{description}"

    def _build_definition(self) -> HermesToolDefinition:
        return HermesToolDefinition(
            name=self.name,
            description=self.description,
            parameters={{
{param_defs}            }},
        )

    async def execute(self, {param_args}**kwargs) -> Dict[str, Any]:
        """执行工具逻辑"""
        logger.info("Executing tool", tool=self.name)
        llm = get_llm_service()
        response = await llm.chat_completion(
            messages=[LLMMessage(role="user", content=f"{description}")],
            temperature=0.3,
        )
        return {{"result": response.content, "tool": self.name}}
'''
        return {
            "success": True,
            "name": name,
            "code": code,
            "message": f"LLM 生成失败({str(e)[:50]})，已使用模板代码",
        }


@router.post("/agents/{agent_id}/skills/chat-generate")
async def chat_generate_skill(agent_id: str, body: dict):
    """通过 LLM 对话生成技能定义"""
    _require_agent_profile(agent_id)
    name = body.get("name", "").strip()
    description = body.get("description", "").strip()

    if not description:
        raise HTTPException(status_code=400, detail="技能描述不能为空")

    # 使用 LLM 生成技能元数据
    from src.core.llm_client import get_llm_service, LLMMessage as _LLMMsg

    llm = get_llm_service()

    prompt = f"""你是一个专利代理领域的 AI Agent 技能设计专家。

请根据以下信息设计一个 Agent 技能：
- 技能名称: {name or '（请根据描述自动命名）'}
- 技能描述: {description}

请输出 JSON 格式（不要输出其他内容）：
{{
  "name": "技能名称（2-6个中文字）",
  "description": "技能的详细功能描述（一句话概括该技能能做什么）",
  "proficiency": 0.85,
  "keywords": ["关键词1", "关键词2", "关键词3", "keyword_en"]
}}

要求：
1. name 简洁有力
2. description 明确描述技能的作用
3. proficiency 在 0.7-0.95 之间
4. keywords 包含 3-6 个关键词（中英文混合），用于任务匹配
"""

    try:
        response = await llm.chat_completion(
            messages=[_LLMMsg(role="user", content=prompt)],
            temperature=0.4,
        )
        content = response.content or ""

        # 解析 JSON
        import json as _json_parse
        # 清理 markdown 包裹
        if "```json" in content:
            content = content.split("```json", 1)[1].split("```", 1)[0]
        elif "```" in content:
            content = content.split("```", 1)[1].split("```", 1)[0]

        skill_data = _json_parse.loads(content.strip())

        # 确保字段完整
        skill_data.setdefault("name", name or "auto_skill")
        skill_data.setdefault("description", description)
        skill_data.setdefault("proficiency", 0.85)
        skill_data.setdefault("keywords", [])

        # 生成 Markdown 内容
        generated_content = f"""# {skill_data['name']}

## 概述
{skill_data['description']}

## 配置
- 名称: {skill_data['name']}
- 熟练度: {skill_data['proficiency']}
- 关键词: {', '.join(skill_data['keywords'])}

## 工作原理
该技能将被注入到 Agent 的 system_prompt 中：
1. 技能描述告诉 LLM 该 Agent 具备此能力
2. 关键词用于任务路由时匹配
3. 熟练度影响 Agent 在该领域的置信度

## 使用场景
当用户输入包含以下关键词时，Agent 会自动运用该技能：
{chr(10).join(f'- {kw}' for kw in skill_data['keywords'])}
"""

        return {
            "success": True,
            "name": skill_data["name"],
            "skill_data": skill_data,
            "generated_content": generated_content,
            "message": "技能已通过 LLM 生成，请确认后添加",
        }
    except Exception as e:
        # 降级：手动生成
        keywords = [w.strip() for w in description.replace(",", " ").split() if len(w.strip()) > 1][:5]
        skill_data = {
            "name": name or "auto_skill",
            "description": description,
            "proficiency": 0.85,
            "keywords": keywords,
        }
        generated_content = f"# {skill_data['name']}\n\n{description}\n\n关键词: {', '.join(keywords)}"

        return {
            "success": True,
            "name": skill_data["name"],
            "skill_data": skill_data,
            "generated_content": generated_content,
            "message": f"LLM 生成失败({str(e)[:50]})，已使用基础模板",
        }


@router.post("/agents/{agent_id}/tools/validate")
async def validate_agent_tool(agent_id: str, body: dict):
    """验证Hermes工具代码语法"""
    _require_agent_profile(agent_id)
    code = body.get("code", "")
    if not code.strip():
        return {"valid": False, "name": None, "error": "代码为空"}

    # Python语法检查
    try:
        compile(code, "<hermes_tool>", "exec")
    except SyntaxError as e:
        return {"valid": False, "name": None, "error": str(e)}

    # 提取函数名
    import ast
    try:
        tree = ast.parse(code)
        func_name = None
        for node in ast.walk(tree):
            if isinstance(node, (ast.AsyncFunctionDef, ast.FunctionDef)):
                func_name = node.name
                break
        return {"valid": True, "name": func_name or "run", "error": None}
    except SyntaxError as e:
        return {"valid": False, "name": None, "error": str(e)}


@router.post("/agents/{agent_id}/tools/hot-plug")
async def hot_plug_agent_tool(agent_id: str, body: dict):
    """热插拔注册新的 Hermes 工具（持久化到 override store + 写入文件）"""
    profile = _require_agent_profile(agent_id)
    name = body.get("name", "").strip()
    description = body.get("description", "").strip()
    code = body.get("code", "")

    if not name:
        raise HTTPException(status_code=400, detail="工具名称不能为空")
    if not code:
        raise HTTPException(status_code=400, detail="工具代码不能为空")

    # 验证代码语法
    try:
        compile(code, "<hermes_tool>", "exec")
    except SyntaxError as e:
        raise HTTPException(status_code=400, detail=f"代码语法错误: {e}")

    # 写入工具文件
    _project_root = _os.path.dirname(
        _os.path.dirname(_os.path.dirname(_os.path.dirname(_os.path.abspath(__file__))))
    )
    tool_file_path = _os.path.join(
        _project_root, "backend", "src", "agents", "hermes", "tools", f"{name}.py"
    )
    try:
        with open(tool_file_path, "w", encoding="utf-8") as f:
            f.write(code)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"写入工具文件失败: {e}")

    # 持久化到 override store（添加到工具列表）
    tool_entry = {
        "id": name,
        "name": name,
        "description": description,
        "enabled": True,
        "category": "analysis",
        "source_code": code,
        "is_hermes": True,
        "config": {},
    }
    _override_store.add_tool(agent_id, tool_entry)

    # 更新工具文件映射
    _KNOWN_TOOL_IMPL_FILES[name] = f"backend/src/agents/hermes/tools/{name}.py"

    return {
        "success": True,
        "name": name,
        "message": f"工具 {name} 已创建并注册（文件: tools/{name}.py）",
    }


@router.get("/agents/{agent_id}/browse")
async def browse_agent_directory(agent_id: str, path: str = ""):
    """浏览Agent工作目录（支持子目录导航）"""
    profile = _require_agent_profile(agent_id)

    _project_root = _os.path.dirname(
        _os.path.dirname(_os.path.dirname(_os.path.dirname(_os.path.abspath(__file__))))
    )

    working_dir = f"./workspace/{profile.profile_id.replace('.', '-')}"
    abs_base = _os.path.normpath(_os.path.join(_project_root, working_dir))

    if not _os.path.isdir(abs_base):
        return {
            "path": path or "/",
            "absolute_path": abs_base,
            "entries": [],
        }

    clean_path = _os.path.normpath(_os.path.join(abs_base, path.lstrip("/")))
    if not clean_path.startswith(abs_base):
        raise HTTPException(status_code=403, detail="路径越权访问")

    if not _os.path.isdir(clean_path):
        raise HTTPException(status_code=404, detail="目录不存在")

    entries = []
    try:
        for name in sorted(_os.listdir(clean_path)):
            full = _os.path.join(clean_path, name)
            is_dir = _os.path.isdir(full)
            size = _os.path.getsize(full) if not is_dir else 0
            entries.append({
                "name": name,
                "path": _os.path.relpath(full, abs_base),
                "type": "directory" if is_dir else "file",
                "size": size,
            })
    except PermissionError as e:
        raise HTTPException(status_code=403, detail=f"无权限访问目录: {e}")

    entries.sort(key=lambda e: (0 if e["type"] == "directory" else 1, e["name"].lower()))
    return {
        "path": path or "/",
        "absolute_path": abs_base,
        "entries": entries,
    }


@router.get("/agents/{agent_id}/file")
async def read_agent_file(agent_id: str, path: str = Query(...)):
    """读取Agent工作目录下的文件内容"""
    profile = _require_agent_profile(agent_id)

    _project_root = _os.path.dirname(
        _os.path.dirname(_os.path.dirname(_os.path.dirname(_os.path.abspath(__file__))))
    )

    working_dir = f"./workspace/{profile.profile_id.replace('.', '-')}"
    abs_base = _os.path.normpath(_os.path.join(_project_root, working_dir))

    clean_path = _os.path.normpath(_os.path.join(abs_base, path.lstrip("/")))
    if not clean_path.startswith(abs_base):
        raise HTTPException(status_code=403, detail="路径越权访问")

    if not _os.path.isfile(clean_path):
        raise HTTPException(status_code=404, detail="文件不存在")

    try:
        import codecs

        for enc in ("utf-8", "utf-8-sig", "gbk", "latin-1"):
            try:
                with codecs.open(clean_path, "r", encoding=enc) as f:
                    content = f.read()
                return {"path": path, "content": content}
            except (UnicodeDecodeError, UnicodeError):
                continue

        # Binary fallback — read as base64
        with open(clean_path, "rb") as f:
            import base64

            raw = f.read()
        return {"path": path, "content": base64.b64encode(raw).decode("ascii"), "encoding": "base64"}
    except PermissionError as e:
        raise HTTPException(status_code=403, detail=f"无权限读取文件: {e}")
    except OSError as e:
        raise HTTPException(status_code=500, detail=f"读取文件失败: {e}")


@router.get("/export/{task_id}")
async def export_patent_docx(task_id: str):
    """导出专利文档为DOCX格式并下载"""
    async with workflow_lock:
        task = tasks_store.get(task_id)
    if not task:
        raise HTTPException(status_code=404, detail="任务不存在")

    patent_data: dict | None = None
    if task.final_patent:
        fp = task.final_patent
        if isinstance(fp, dict):
            patent_data = fp.get("patent_draft") or fp
        elif fp.patent_draft:
            patent_data = fp.patent_draft.model_dump()
    elif task.draft_doc:
        patent_data = task.draft_doc.model_dump()

    if not patent_data:
        raise HTTPException(status_code=400, detail="该任务尚无专利文档可导出")

    from src.document_gen.generator import generate_patent_docx
    filepath = await asyncio.to_thread(generate_patent_docx, patent_data, task_id)

    return FileResponse(
        path=filepath,
        filename=f"patent_{task_id}.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@router.get("/workflows/{task_id}/export/docx")
async def export_workflow_patent_docx(task_id: str):
    """从工作流导出专利文档为DOCX格式"""
    async with workflow_lock:
        context = workflow_engine.get_workflow(task_id)
    if not context:
        raise HTTPException(status_code=404, detail="工作流不存在")

    patent_draft = context.patent_draft
    if not patent_draft:
        raise HTTPException(status_code=400, detail="工作流尚未生成专利文档，请等待撰写阶段完成")

    # 将工作流输出格式转换为 generate_patent_docx 需要的格式
    draft_data = patent_draft if isinstance(patent_draft, dict) else patent_draft

    # 工作流输出格式: {claims: {independent_claim, dependent_claims}, description: {...}, abstract}
    # generate_patent_docx 需要的格式: PatentDraft model 或兼容 dict
    # 做适配转换
    claims_raw = draft_data.get("claims", {})
    desc_raw = draft_data.get("description", {})

    adapted_data = {
        "title": context.requirement_analysis.get("patent_title", "专利申请文件"),
        "technical_field": desc_raw.get("technical_field", ""),
        "background_art": {"section_name": "背景技术", "content": desc_raw.get("background_art", ""), "word_count": len(desc_raw.get("background_art", ""))},
        "summary_of_invention": {"section_name": "发明内容", "content": desc_raw.get("summary_of_invention", ""), "word_count": len(desc_raw.get("summary_of_invention", ""))},
        "description_of_drawings": {"section_name": "附图说明", "content": desc_raw.get("description_of_drawings", ""), "word_count": len(desc_raw.get("description_of_drawings", ""))} if desc_raw.get("description_of_drawings") else None,
        "detailed_description": {"section_name": "具体实施方式", "content": desc_raw.get("detailed_description", ""), "word_count": len(desc_raw.get("detailed_description", ""))},
        "claims": [],
        "abstract": draft_data.get("abstract", ""),
        "figures": draft_data.get("figures", []),
    }

    # 转换权利要求
    import re
    def _strip_claim_prefix(text: str) -> str:
        """去除LLM生成的权利要求编号前缀，如 '1. ', '2. 根据权利要求1所述的方法，'"""
        # 去掉开头的数字编号 (如 "1. ", "2. ", "10. ")
        text = re.sub(r'^\d+\.\s*', '', text.strip())
        return text

    if isinstance(claims_raw, dict):
        ind_claim = claims_raw.get("independent_claim", "")
        if ind_claim:
            adapted_data["claims"].append({"claim_number": 1, "claim_type": "independent", "content": _strip_claim_prefix(ind_claim), "dependencies": []})
        for i, dep in enumerate(claims_raw.get("dependent_claims", []), 2):
            # 从属权利要求：提取实际依赖的权利要求号
            dep_text = _strip_claim_prefix(dep)
            dep_nums = re.findall(r'根据权利要求(\d+)', dep_text)
            dependencies = [int(n) for n in dep_nums] if dep_nums else [1]
            adapted_data["claims"].append({"claim_number": i, "claim_type": "dependent", "content": dep_text, "dependencies": dependencies})

    from src.document_gen.generator import generate_patent_docx
    # 清除旧的导出文件以确保重新生成
    from pathlib import Path as _Path
    _export_dir = _Path("./exports") / task_id
    if _export_dir.exists():
        import shutil
        shutil.rmtree(_export_dir)
    try:
        filepath = await asyncio.to_thread(generate_patent_docx, adapted_data, task_id)
    except Exception as e:
        # 如果 PatentDraft 验证失败，直接用简单 docx 生成
        import traceback
        with open("/tmp/docx_export_error.log", "w") as f:
            f.write(f"{type(e).__name__}: {e}\n")
            traceback.print_exc(file=f)
        from pathlib import Path
        from docx import Document

        doc = Document()
        doc.add_heading(adapted_data["title"], level=0)

        doc.add_heading("权利要求书", level=1)
        if isinstance(claims_raw, dict):
            doc.add_paragraph(claims_raw.get("independent_claim", ""))
            for dep in claims_raw.get("dependent_claims", []):
                doc.add_paragraph(dep)

        doc.add_heading("说明书", level=1)
        doc.add_heading("【技术领域】", level=2)
        doc.add_paragraph(desc_raw.get("technical_field", ""))
        doc.add_heading("【背景技术】", level=2)
        doc.add_paragraph(desc_raw.get("background_art", ""))
        doc.add_heading("【发明内容】", level=2)
        doc.add_paragraph(desc_raw.get("summary_of_invention", ""))
        doc.add_heading("【具体实施方式】", level=2)
        doc.add_paragraph(desc_raw.get("detailed_description", ""))

        doc.add_heading("说明书摘要", level=1)
        doc.add_paragraph(draft_data.get("abstract", ""))

        export_dir = Path("./exports") / task_id
        export_dir.mkdir(parents=True, exist_ok=True)
        filepath = str(export_dir / f"{task_id}_专利申请书.docx")
        doc.save(filepath)

    return FileResponse(
        path=filepath,
        filename=f"patent_{task_id}.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@router.post("/workflows/{task_id}/pause")
async def pause_workflow(task_id: str):
    """暂停工作流执行"""
    async with workflow_lock:
        context = workflow_engine.get_workflow(task_id)
    if not context:
        raise HTTPException(status_code=404, detail="工作流不存在")

    context.is_paused = True
    _append_workflow_event(
        task_id=task_id,
        agent="workflow_engine",
        message="工作流已暂停",
        event_type="workflow.paused",
    )
    await _persist_events(task_id)
    return {"task_id": task_id, "status": "paused"}


@router.post("/workflows/{task_id}/resume")
async def resume_workflow(task_id: str):
    """恢复暂停的工作流"""
    async with workflow_lock:
        context = workflow_engine.get_workflow(task_id)
    if not context:
        raise HTTPException(status_code=404, detail="工作流不存在")

    context.is_paused = False
    _append_workflow_event(
        task_id=task_id,
        agent="workflow_engine",
        message="工作流已恢复",
        event_type="workflow.resumed",
    )
    await _persist_events(task_id)
    return {"task_id": task_id, "status": "resumed"}
