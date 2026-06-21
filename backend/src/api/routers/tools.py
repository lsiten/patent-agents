import os
from pathlib import Path
from typing import List, Dict, Any, Optional
from fastapi import APIRouter, HTTPException, UploadFile, File
from loguru import logger
import yaml

router = APIRouter(tags=["tools"])

_TOOLS_AGENT_PROFILES_ROOT = Path(__file__).resolve().parents[3] / "hermes_home" / "profiles" / "tools_agent"


def _load_tool_skill(skill_name: str) -> Optional[Dict[str, Any]]:
    """加载单个工具的 SKILL.md 文件内容"""
    skill_dir = _TOOLS_AGENT_PROFILES_ROOT / "skills" / skill_name
    skill_file = skill_dir / "SKILL.md"
    
    if not skill_file.exists():
        logger.warning(f"Skill file not found: {skill_file}")
        return None
    
    try:
        content = skill_file.read_text(encoding="utf-8")
        lines = content.split("\n")
        
        yaml_metadata = {}
        if lines and lines[0] == "---":
            end_idx = None
            for i in range(1, len(lines)):
                if lines[i] == "---":
                    end_idx = i
                    break
            if end_idx:
                yaml_content = "\n".join(lines[1:end_idx])
                yaml_metadata = yaml.safe_load(yaml_content)
        
        # 提取基本信息
        name = yaml_metadata.pop("name", skill_name)
        display_name = yaml_metadata.pop("display_name", "")
        description = yaml_metadata.pop("description", "")
        
        # 返回前端期望的格式
        result = {
            "name": name,
            "display_name": display_name,
            "description": description,
            "metadata": yaml_metadata,
            "content": content,
        }
        
        return result
    except Exception as e:
        logger.error(f"Failed to load skill {skill_name}: {e}")
        return None


def _load_all_tools() -> List[Dict[str, Any]]:
    """加载所有工具列表"""
    skills_dir = _TOOLS_AGENT_PROFILES_ROOT / "skills"
    
    if not skills_dir.exists():
        logger.warning(f"Skills directory not found: {skills_dir}")
        return []
    
    tools = []
    for skill_dir in sorted(skills_dir.iterdir()):
        if skill_dir.is_dir():
            skill_name = skill_dir.name
            skill_data = _load_tool_skill(skill_name)
            if skill_data:
                tools.append(skill_data)
    
    return tools


@router.get("/tools", response_model=List[Dict[str, Any]])
async def get_tools_list():
    """获取所有工具列表"""
    tools = _load_all_tools()
    return tools


@router.get("/tools/{tool_name}", response_model=Dict[str, Any])
async def get_tool_detail(tool_name: str):
    """获取单个工具的详细信息"""
    tool_data = _load_tool_skill(tool_name)
    
    if not tool_data:
        raise HTTPException(status_code=404, detail=f"工具 {tool_name} 不存在")
    
    return tool_data


@router.get("/tools/categories", response_model=List[str])
async def get_tool_categories():
    """获取所有工具分类"""
    tools = _load_all_tools()
    categories = set()
    
    for tool in tools:
        category = tool.get("metadata", {}).get("category")
        if category:
            categories.add(category)
    
    return sorted(list(categories))


@router.get("/tools/category/{category}", response_model=List[Dict[str, Any]])
async def get_tools_by_category(category: str):
    """按分类获取工具列表"""
    tools = _load_all_tools()
    filtered_tools = [
        tool for tool in tools 
        if tool.get("metadata", {}).get("category") == category
    ]
    
    return filtered_tools


async def _extract_text_from_file(file: UploadFile) -> str:
    """从上传文件中提取文本内容"""
    filename = (file.filename or "uploaded").strip()
    suffix = os.path.splitext(filename)[1].lower()

    raw_bytes = await file.read()
    if not raw_bytes:
        raise HTTPException(status_code=400, detail="上传文件为空")

    max_bytes = 10 * 1024 * 1024
    if len(raw_bytes) > max_bytes:
        raise HTTPException(
            status_code=413,
            detail=f"文件过大（{len(raw_bytes) // 1024} KB），上限 {max_bytes // 1024 // 1024} MB",
        )

    extracted_text = ""

    try:
        if suffix == ".txt" or suffix == ".md" or file.content_type == "text/plain":
            try:
                extracted_text = raw_bytes.decode("utf-8")
            except UnicodeDecodeError:
                extracted_text = raw_bytes.decode("gb18030", errors="replace")
        elif suffix == ".docx":
            try:
                from docx import Document
            except ImportError as e:
                raise HTTPException(
                    status_code=500,
                    detail="服务器未安装 python-docx，无法解析 docx 文件",
                ) from e

            import io
            doc = Document(io.BytesIO(raw_bytes))
            paragraphs = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            paragraphs.append(cell_text)
            extracted_text = "\n".join(paragraphs).strip()
        elif suffix == ".pdf":
            try:
                import fitz
            except ImportError as e:
                raise HTTPException(
                    status_code=500,
                    detail="服务器未安装 PyMuPDF，无法解析 pdf 文件",
                ) from e

            import io
            doc = fitz.open(stream=raw_bytes, filetype="pdf")
            try:
                pages = []
                for page_num in range(len(doc)):
                    text = doc[page_num].get_text().strip()
                    if text:
                        pages.append(text)
                extracted_text = "\n\n".join(pages).strip()
            finally:
                doc.close()
        else:
            raise HTTPException(
                status_code=400,
                detail=f"不支持的文件类型: {suffix or file.content_type}，仅支持 .txt、.md、.docx 和 .pdf",
            )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"文件解析失败: {e}")
        raise HTTPException(status_code=500, detail=f"文件解析失败: {str(e)}")

    return extracted_text


@router.post("/tools/upload")
async def upload_tool_file(file: UploadFile = File(...)):
    """上传文件并提取文本内容"""
    try:
        content = await _extract_text_from_file(file)
        return {"content": content, "filename": file.filename}
    except Exception as e:
        logger.error(f"File upload error: {e}")
        raise HTTPException(status_code=500, detail=f"文件处理失败: {str(e)}")