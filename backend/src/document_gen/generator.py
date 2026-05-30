"""
Patent document generator — produces professionally formatted .docx files
based on an extracted PatentFeatureProfile (not direct file access).

The generator reads a persisted feature profile (JSON) that describes the exact
formatting characteristics of finalized patents. It never reads the original
docx reference files at generation time.

Flow:
  1. Feature profile extracted once from 定稿文件/ → saved as JSON
  2. Generator loads the JSON profile at startup
  3. All formatting decisions are driven by the profile
"""

from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from src.document_gen.feature_profile import (
    PatentFeatureProfile,
    load_profile,
)
from src.models.domain import (
    Claim,
    FinalPatent,
    PatentDraft,
    PatentTask,
)

EXPORT_DIR = Path(__file__).resolve().parent.parent.parent / "exports"

# Cached profile (loaded once from persisted JSON or defaults)
_cached_profile: Optional[PatentFeatureProfile] = None


def _get_profile() -> PatentFeatureProfile:
    """Load the feature profile (cached after first call)."""
    global _cached_profile
    if _cached_profile is None:
        _cached_profile = load_profile()
    return _cached_profile


def set_profile(profile: PatentFeatureProfile) -> None:
    """Override the cached profile (for testing or custom profiles)."""
    global _cached_profile
    _cached_profile = profile


def reset_profile() -> None:
    """Reset cached profile (forces reload on next use)."""
    global _cached_profile
    _cached_profile = None


# ─── Low-level paragraph helpers ─────────────────────────────────────────────


def _set_line_spacing_exact(para, emu: int) -> None:
    """Set exact line spacing via XML. EMU → twips for w:line attribute."""
    # w:line uses twips (1/20 of a point). EMU / 635 = twips
    twips = round(emu / 635)
    pPr = para._element.get_or_add_pPr()
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        pPr.append(spacing)
    spacing.set(qn("w:line"), str(twips))
    spacing.set(qn("w:lineRule"), "exact")


def _set_run_font(run, font_name: str, size_pt: float, bold: bool = False):
    """Set font properties on a run with proper CJK handling."""
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold
    # Set East-Asian, ASCII, and HAnsi fonts for CJK
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)


def _get_alignment(align_str: str):
    """Convert alignment string to python-docx constant."""
    mapping = {
        "JUSTIFY": WD_ALIGN_PARAGRAPH.JUSTIFY,
        "CENTER": WD_ALIGN_PARAGRAPH.CENTER,
        "LEFT": WD_ALIGN_PARAGRAPH.LEFT,
        "RIGHT": WD_ALIGN_PARAGRAPH.RIGHT,
    }
    return mapping.get(align_str.upper(), WD_ALIGN_PARAGRAPH.JUSTIFY)


def add_body_paragraph(
    doc: Document,
    text: str,
    profile: PatentFeatureProfile,
    first_line_indent: bool = True,
    bold: bool = False,
) -> None:
    """Add a body text paragraph formatted according to the feature profile."""
    para = doc.add_paragraph()
    run = para.add_run(text)

    typo = profile.typography
    pfmt = profile.paragraph_format

    _set_run_font(run, typo.body_font_name, typo.body_font_size_pt, bold)

    # Alignment
    para.paragraph_format.alignment = _get_alignment(pfmt.body_alignment)

    # First line indent
    if first_line_indent:
        para.paragraph_format.first_line_indent = Cm(pfmt.body_first_line_indent_cm)

    # Line spacing
    _set_line_spacing_exact(para, pfmt.body_line_spacing_emu)


def add_section_heading(
    doc: Document,
    text: str,
    profile: PatentFeatureProfile,
) -> None:
    """Add a description sub-section heading (e.g. 技术领域, 背景技术)."""
    para = doc.add_paragraph()
    run = para.add_run(text)

    typo = profile.typography
    pfmt = profile.paragraph_format

    _set_run_font(
        run,
        typo.section_heading_font_name,
        typo.section_heading_font_size_pt,
        bold=typo.section_heading_bold,
    )

    para.paragraph_format.alignment = _get_alignment(pfmt.section_heading_alignment)
    _set_line_spacing_exact(para, pfmt.body_line_spacing_emu)


def add_document_heading(
    doc: Document,
    text: str,
    profile: PatentFeatureProfile,
) -> None:
    """Add a document-level heading (e.g. 权利要求书, 说明书)."""
    struct = profile.document_structure
    pfmt = profile.paragraph_format
    typo = profile.typography

    # Space characters apart using separator from profile
    spaced_text = struct.heading_char_separator.join(list(text))

    h = doc.add_heading(spaced_text, level=1)

    # Apply font to all runs
    for run in h.runs:
        _set_run_font(
            run, typo.doc_heading_font_name, typo.section_heading_font_size_pt,
            bold=typo.doc_heading_bold,
        )

    h.alignment = _get_alignment(pfmt.doc_heading_alignment)


# ─── Section/Page Management ────────────────────────────────────────────────


def _set_section_margins(section, margins: tuple) -> None:
    """Set section margins from (top, bottom, left, right) in cm."""
    top, bottom, left, right = margins
    section.top_margin = Cm(top)
    section.bottom_margin = Cm(bottom)
    section.left_margin = Cm(left)
    section.right_margin = Cm(right)


def _set_page_size(section, profile: PatentFeatureProfile) -> None:
    """Set page size from profile."""
    section.page_width = Cm(profile.page_layout.page_width_cm)
    section.page_height = Cm(profile.page_layout.page_height_cm)


def _get_margins_for_section(section_name: str, profile: PatentFeatureProfile) -> tuple:
    """Get margins tuple for a named section from profile."""
    layout = profile.page_layout
    mapping = {
        "摘要": layout.margins_abstract,
        "摘要附图": layout.margins_abstract_drawings,
        "权利要求书": layout.margins_claims,
        "说明书": layout.margins_description,
        "说明书附图": layout.margins_description_drawings,
    }
    return mapping.get(section_name, layout.margins_abstract)


def add_new_section(
    doc: Document,
    section_name: str,
    profile: PatentFeatureProfile,
) -> None:
    """Add a new section with page break and margins from profile."""
    doc.add_section()
    new_section = doc.sections[-1]
    _set_page_size(new_section, profile)
    margins = _get_margins_for_section(section_name, profile)
    _set_section_margins(new_section, margins)


# ─── Document Building ───────────────────────────────────────────────────────


def build_abstract_section(doc: Document, draft: PatentDraft, profile: PatentFeatureProfile) -> None:
    """Generate 说明书摘要 section."""
    add_document_heading(doc, "说明书摘要", profile)
    add_body_paragraph(doc, draft.abstract, profile)


def build_abstract_drawings_section(doc: Document, draft: PatentDraft, profile: PatentFeatureProfile, figure_paths: Optional[List[Dict]] = None) -> None:
    """Generate 摘要附图 section."""
    add_new_section(doc, "摘要附图", profile)
    add_document_heading(doc, "摘要附图", profile)
    if figure_paths and len(figure_paths) > 0:
        # 插入第一张图作为摘要附图
        from docx.shared import Inches
        try:
            doc.add_picture(figure_paths[0]["path"], width=Inches(5.5))
        except Exception:
            add_body_paragraph(doc, "（附图生成失败）", profile, first_line_indent=False)
    else:
        add_body_paragraph(doc, "（无附图）", profile, first_line_indent=False)


def build_claims_section(doc: Document, draft: PatentDraft, profile: PatentFeatureProfile) -> None:
    """Generate 权利要求书 section."""
    add_new_section(doc, "权利要求书", profile)
    add_document_heading(doc, "权利要求书", profile)

    if not draft.claims:
        add_body_paragraph(doc, "（无权利要求内容）", profile)
        return

    for claim in draft.claims:
        text = _format_claim(claim, profile)
        add_body_paragraph(doc, text, profile)


def _format_claim(claim: Claim, profile: PatentFeatureProfile) -> str:
    """Format a claim using content patterns from profile."""
    sep = profile.content_patterns.claim_numbering_separator
    prefix = f"{claim.claim_number}{sep}"
    if claim.dependencies:
        dep_ref = "、".join(str(d) for d in claim.dependencies)
        dep_fmt = profile.content_patterns.claim_dependent_format.replace("{n}", dep_ref)
        prefix = f"{claim.claim_number}{sep}{dep_fmt}"
    return f"{prefix}{claim.content}"


def build_description_section(doc: Document, draft: PatentDraft, profile: PatentFeatureProfile) -> None:
    """Generate 说明书 section with all sub-sections."""
    add_new_section(doc, "说明书", profile)
    add_document_heading(doc, "说明书", profile)

    # 技术领域
    add_section_heading(doc, "技术领域", profile)
    add_body_paragraph(doc, draft.technical_field, profile)

    # 背景技术
    add_section_heading(doc, "背景技术", profile)
    _add_multiline_content(doc, draft.background_art.content, profile)

    # 发明内容
    add_section_heading(doc, "发明内容", profile)
    _add_multiline_content(doc, draft.summary_of_invention.content, profile)

    # 附图说明 (optional)
    if draft.description_of_drawings:
        add_section_heading(doc, "附图说明", profile)
        _add_multiline_content(doc, draft.description_of_drawings.content, profile)

    # 具体实施方式
    add_section_heading(doc, "具体实施方式", profile)
    _add_multiline_content(doc, draft.detailed_description.content, profile)


def build_description_drawings_section(
    doc: Document, draft: PatentDraft, profile: PatentFeatureProfile, figure_paths: Optional[List[Dict]] = None
) -> None:
    """Generate 说明书附图 section."""
    add_new_section(doc, "说明书附图", profile)
    add_document_heading(doc, "说明书附图", profile)
    if figure_paths and len(figure_paths) > 0:
        from docx.shared import Inches
        for fig_info in figure_paths:
            try:
                # 添加图标题
                fig_title = f"图{fig_info.get('figure_number', '')}: {fig_info.get('title', '')}"
                add_body_paragraph(doc, fig_title, profile, first_line_indent=False, bold=True)
                # 插入图片
                doc.add_picture(fig_info["path"], width=Inches(5.5))
                # 空行
                doc.add_paragraph("")
            except Exception:
                add_body_paragraph(doc, f"（图{fig_info.get('figure_number', '')}生成失败）", profile, first_line_indent=False)
    else:
        add_body_paragraph(doc, "（无附图）", profile, first_line_indent=False)


def _add_multiline_content(doc: Document, content: str, profile: PatentFeatureProfile) -> None:
    """Add content that may contain multiple paragraphs. Strips Markdown syntax."""
    content = _strip_markdown(content)
    paragraphs = [p.strip() for p in content.split("\n") if p.strip()]
    for para_text in paragraphs:
        add_body_paragraph(doc, para_text, profile)


def _strip_markdown(text: str) -> str:
    """Remove Markdown formatting from LLM-generated text for clean docx output.
    
    Handles: headings (#), bold (**), italic (*), code (`), lists (- / *), 
    numbered lists, links [text](url), horizontal rules (---).
    """
    import re
    if not text:
        return text

    lines = text.split("\n")
    cleaned_lines = []

    for line in lines:
        # Skip pure horizontal rules
        if re.match(r'^[\s]*[-*_]{3,}\s*$', line):
            continue

        # Remove heading markers (# ## ### etc.)
        line = re.sub(r'^#{1,6}\s+', '', line)

        # Remove bold markers (**text** or __text__)
        line = re.sub(r'\*\*(.+?)\*\*', r'\1', line)
        line = re.sub(r'__(.+?)__', r'\1', line)

        # Remove italic markers (*text* or _text_) — careful not to break underscores in words
        line = re.sub(r'(?<!\w)\*([^*]+?)\*(?!\w)', r'\1', line)
        line = re.sub(r'(?<!\w)_([^_]+?)_(?!\w)', r'\1', line)

        # Remove inline code backticks
        line = re.sub(r'`([^`]+?)`', r'\1', line)

        # Remove code block markers
        if re.match(r'^```', line):
            continue

        # Convert list markers to proper paragraph text
        # - item  or  * item  → item
        line = re.sub(r'^\s*[-*+]\s+', '', line)
        # 1. item  or  1) item  → item (numbered list — keep number for patent claims context)
        # Don't strip numbered lists as they might be intentional in patent context

        # Remove link syntax [text](url) → text
        line = re.sub(r'\[([^\]]+?)\]\([^)]+?\)', r'\1', line)

        # Remove image syntax ![alt](url)
        line = re.sub(r'!\[([^\]]*?)\]\([^)]+?\)', r'\1', line)

        # Remove blockquote markers
        line = re.sub(r'^>\s*', '', line)

        cleaned_lines.append(line)

    return "\n".join(cleaned_lines)


# ─── Main Entry Point ────────────────────────────────────────────────────────


def generate_patent_docx(
    patent_data: Dict[str, Any] | PatentTask | FinalPatent | PatentDraft,
    task_id_or_patent: str | FinalPatent | None = None,
    output_dir: Optional[Path] = None,
) -> str:
    """
    Generate a professional Chinese patent application .docx file
    based on the extracted feature profile.

    The feature profile (PatentFeatureProfile) describes all formatting:
    - Page layout, margins per section
    - Typography (font, size, bold)
    - Paragraph format (indent, line spacing, alignment)
    - Document structure (sections, headings)
    - Content patterns (claim format, numbering)

    The generator never reads the original reference docx files directly.

    Accepts data in multiple forms:
    - (dict, task_id)           — raw dict + task ID string
    - (PatentTask, FinalPatent) — model objects from workflow engine
    - (PatentDraft, task_id)    — draft model + task ID string
    - (FinalPatent,)            — FinalPatent model object

    Returns:
        Absolute path to the generated .docx file.
    """
    # Normalize inputs to (draft_dict, task_id)
    figure_paths: Optional[List[Dict]] = None
    if isinstance(patent_data, dict):
        draft_dict: dict = patent_data
        task_id: str = task_id_or_patent  # type: ignore[assignment]
        # 提取附图路径（如果有）
        figure_paths = patent_data.get("figures") if isinstance(patent_data.get("figures"), list) else None
    elif isinstance(patent_data, FinalPatent):
        draft_dict = patent_data.patent_draft.model_dump()
        task_id = patent_data.task_id
    elif isinstance(patent_data, PatentTask):
        fp = task_id_or_patent
        assert isinstance(fp, FinalPatent), "PatentTask needs FinalPatent as second arg"
        draft_dict = fp.patent_draft.model_dump()
        task_id = patent_data.task_id
    elif isinstance(patent_data, PatentDraft):
        draft_dict = patent_data.model_dump()
        task_id = task_id_or_patent  # type: ignore[assignment]
    else:
        raise TypeError(f"Unsupported patent_data type: {type(patent_data)}")

    # Parse draft dict through Pydantic for validation
    draft = PatentDraft(**draft_dict)

    if output_dir is None:
        output_dir = EXPORT_DIR / task_id
    Path(output_dir).mkdir(parents=True, exist_ok=True)

    file_path = output_dir / f"{task_id}_专利申请书.docx"

    # Load feature profile (from persisted JSON or defaults)
    profile = _get_profile()

    # Create document
    doc = Document()

    # Configure first section (摘要)
    first_section = doc.sections[0]
    _set_page_size(first_section, profile)
    margins = _get_margins_for_section("摘要", profile)
    _set_section_margins(first_section, margins)

    # Build all 5 sections per document structure
    build_abstract_section(doc, draft, profile)
    build_abstract_drawings_section(doc, draft, profile, figure_paths)
    build_claims_section(doc, draft, profile)
    build_description_section(doc, draft, profile)
    build_description_drawings_section(doc, draft, profile, figure_paths)

    doc.save(str(file_path))
    return str(file_path)
