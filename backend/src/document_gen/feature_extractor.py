"""
Patent Feature Extractor — analyzes ALL finalized patent B-files
and produces a comprehensive PatentFeatureProfile.

The extraction process:
1. Iterates over all B-*.docx files in the 定稿文件/ directory
2. For each file, extracts formatting and content characteristics
3. Aggregates/averages across all files to produce a single profile
4. Persists the profile as JSON for use by the generator

This decouples the generator from the original docx files:
  extract once → persist → generate from profile
"""

import re
import statistics
from collections import Counter
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Tuple

from docx import Document
from docx.oxml.ns import qn
from loguru import logger

from src.document_gen.feature_profile import (
    ContentPatterns,
    DEFAULT_PROFILE_PATH,
    DocumentStructure,
    PageLayout,
    ParagraphFormat,
    PatentFeatureProfile,
    Typography,
)


# ─── Per-file extraction ─────────────────────────────────────────────────────


def _extract_page_layout(doc: Document) -> Dict:
    """Extract page layout features from a single document."""
    result = {
        "page_width_cm": None,
        "page_height_cm": None,
        "section_margins": [],  # list of (top, bottom, left, right) per section
    }

    if doc.sections:
        s = doc.sections[0]
        result["page_width_cm"] = round(s.page_width.cm, 2)
        result["page_height_cm"] = round(s.page_height.cm, 2)

    for section in doc.sections:
        result["section_margins"].append((
            round(section.top_margin.cm, 2),
            round(section.bottom_margin.cm, 2),
            round(section.left_margin.cm, 2),
            round(section.right_margin.cm, 2),
        ))

    return result


def _extract_typography(doc: Document) -> Dict:
    """Extract font/typography features from a single document."""
    body_fonts = []
    body_sizes = []
    heading_fonts = []
    heading_bold = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text or not para.runs:
            continue

        run = para.runs[0]
        style_name = para.style.name if para.style else ""

        if "Heading" in style_name:
            if run.font.name:
                heading_fonts.append(run.font.name)
        elif "Body Text" in style_name or style_name == "Normal":
            if run.font.name:
                body_fonts.append(run.font.name)
            if run.font.size:
                body_sizes.append(run.font.size.pt)
            # Check for section heading (bold body text)
            if run.font.bold and text in (
                "技术领域", "背景技术", "发明内容", "附图说明", "具体实施方式"
            ):
                heading_bold.append(True)

    return {
        "body_fonts": body_fonts,
        "body_sizes": body_sizes,
        "heading_fonts": heading_fonts,
        "section_heading_bold": heading_bold,
    }


def _extract_paragraph_format(doc: Document) -> Dict:
    """Extract paragraph formatting features from a single document."""
    indents = []
    line_spacings = []
    alignments = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text or not para.runs:
            continue

        style_name = para.style.name if para.style else ""
        if "Heading" in style_name:
            continue

        fmt = para.paragraph_format

        # First line indent
        if fmt.first_line_indent:
            indents.append(round(fmt.first_line_indent.cm, 2))

        # Line spacing from XML (w:line value in twips = 1/20 pt)
        # Convert to EMU: twips * 635 = EMU (for exact line spacing)
        pPr = para._element.find(qn("w:pPr"))
        if pPr is not None:
            spacing = pPr.find(qn("w:spacing"))
            if spacing is not None:
                line_val = spacing.get(qn("w:line"))
                line_rule = spacing.get(qn("w:lineRule"))
                if line_val and line_rule == "exact":
                    try:
                        twips = int(line_val)
                        # Store as EMU for consistency with python-docx
                        emu = twips * 635
                        line_spacings.append(emu)
                    except ValueError:
                        pass

        # Alignment
        if fmt.alignment is not None:
            alignments.append(str(fmt.alignment))

    return {
        "indents": indents,
        "line_spacings": line_spacings,
        "alignments": alignments,
    }


def _extract_document_structure(doc: Document) -> Dict:
    """Extract structural features (headings, sections) from a document."""
    headings_found = []
    description_sections_found = []
    num_sections = len(doc.sections)

    desc_headings = {"技术领域", "背景技术", "发明内容", "附图说明", "具体实施方式"}

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style_name = para.style.name if para.style else ""

        if "Heading" in style_name:
            # Normalize spaced heading
            normalized = re.sub(r"\s+", "", text)
            headings_found.append(normalized)

        # Check description sub-sections
        if text in desc_headings and para.runs and para.runs[0].font.bold:
            description_sections_found.append(text)

    return {
        "headings": headings_found,
        "description_sections": description_sections_found,
        "num_sections": num_sections,
    }


def _extract_content_patterns(doc: Document) -> Dict:
    """Extract content/writing style patterns from a document."""
    claims = []
    current_claim = []
    in_claims = False
    abstract_text = ""
    in_abstract = False
    description_text = ""
    in_description = False

    claim_start = re.compile(r"^\d+[、．.]")

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style_name = para.style.name if para.style else ""
        normalized = re.sub(r"\s+", "", text)

        # Track which section we're in
        if "Heading" in style_name:
            if normalized == "说明书摘要":
                in_abstract = True
                in_claims = False
                in_description = False
            elif normalized == "权利要求书":
                in_claims = True
                in_abstract = False
                in_description = False
            elif normalized == "说明书":
                in_description = True
                in_claims = False
                in_abstract = False
            elif normalized in ("摘要附图", "说明书附图"):
                in_abstract = False
                in_claims = False
                in_description = False
            continue

        if in_abstract:
            abstract_text += text
        elif in_claims:
            if claim_start.match(text):
                if current_claim:
                    claims.append("\n".join(current_claim))
                current_claim = [text]
            else:
                current_claim.append(text)
        elif in_description:
            description_text += text + "\n"

    if current_claim:
        claims.append("\n".join(current_claim))

    # Analyze patterns
    independent_claims = sum(1 for c in claims if "其特征在于" in c and not "根据权利要求" in c)
    claim_lengths = [len(c) for c in claims]

    return {
        "num_claims": len(claims),
        "num_independent_claims": independent_claims,
        "claim_lengths": claim_lengths,
        "abstract_length": len(abstract_text),
        "description_length": len(description_text),
        "has_characteristic_marker": any("其特征在于" in c for c in claims),
        "has_dependent_format": any("根据权利要求" in c for c in claims),
        "abstract_starts_with_benfaming": abstract_text.startswith("本发明"),
        "has_explanation_markers": "需要说明的是" in description_text,
        "has_closing_statement": "最后应说明的是" in description_text,
    }


# ─── Aggregation ─────────────────────────────────────────────────────────────


def _most_common(items: List, default=None):
    """Return most common item in a list."""
    if not items:
        return default
    counter = Counter(items)
    return counter.most_common(1)[0][0]


def _median_or_default(items: List[float], default: float) -> float:
    """Return median or default if list empty."""
    if not items:
        return default
    return round(statistics.median(items), 2)


def extract_profile(docx_dir: Path) -> PatentFeatureProfile:
    """
    Extract a PatentFeatureProfile by analyzing ALL B-files in the given directory.

    Iterates every subdirectory, finds B-*.docx files, extracts features from each,
    and aggregates into a single profile representing the common characteristics.
    """
    if not docx_dir.exists():
        logger.warning(f"定稿文件目录不存在: {docx_dir}")
        return PatentFeatureProfile()

    b_files: List[Path] = []
    for subdir in sorted(docx_dir.iterdir()):
        if not subdir.is_dir():
            continue
        for f in subdir.glob("B-*.docx"):
            b_files.append(f)

    if not b_files:
        logger.warning(f"未找到 B 文件: {docx_dir}")
        return PatentFeatureProfile()

    logger.info(f"开始特征提取，共 {len(b_files)} 个 B 文件")

    # Collect per-file features
    all_layouts = []
    all_typo = []
    all_para = []
    all_structure = []
    all_content = []

    for b_file in b_files:
        try:
            doc = Document(str(b_file))
            all_layouts.append(_extract_page_layout(doc))
            all_typo.append(_extract_typography(doc))
            all_para.append(_extract_paragraph_format(doc))
            all_structure.append(_extract_document_structure(doc))
            all_content.append(_extract_content_patterns(doc))
            logger.debug(f"提取特征: {b_file.name}")
        except Exception as e:
            logger.error(f"提取特征失败 {b_file}: {e}")

    if not all_layouts:
        return PatentFeatureProfile()

    # ─── Aggregate page layout ─────────
    page_widths = [l["page_width_cm"] for l in all_layouts if l["page_width_cm"]]
    page_heights = [l["page_height_cm"] for l in all_layouts if l["page_height_cm"]]

    # For margins, take the most common pattern (they should all be the same)
    # Aggregate per section index
    section_margin_lists = [[], [], [], [], []]  # 5 sections
    for layout in all_layouts:
        for i, m in enumerate(layout["section_margins"][:5]):
            section_margin_lists[i].append(m)

    def _median_margin(margin_list) -> Tuple[float, float, float, float]:
        if not margin_list:
            return (2.60, 2.00, 2.80, 1.80)
        tops = [m[0] for m in margin_list]
        bots = [m[1] for m in margin_list]
        lefts = [m[2] for m in margin_list]
        rights = [m[3] for m in margin_list]
        return (
            round(statistics.median(tops), 2),
            round(statistics.median(bots), 2),
            round(statistics.median(lefts), 2),
            round(statistics.median(rights), 2),
        )

    page_layout = PageLayout(
        page_width_cm=_median_or_default(page_widths, 21.0),
        page_height_cm=_median_or_default(page_heights, 29.7),
        margins_abstract=_median_margin(section_margin_lists[0]),
        margins_abstract_drawings=_median_margin(section_margin_lists[1]),
        margins_claims=_median_margin(section_margin_lists[2]),
        margins_description=_median_margin(section_margin_lists[3]),
        margins_description_drawings=_median_margin(section_margin_lists[4]),
    )

    # ─── Aggregate typography ─────────
    all_body_fonts = []
    all_body_sizes = []
    all_heading_fonts = []

    for t in all_typo:
        all_body_fonts.extend(t["body_fonts"])
        all_body_sizes.extend(t["body_sizes"])
        all_heading_fonts.extend(t["heading_fonts"])

    typography = Typography(
        body_font_name=_most_common(all_body_fonts, "楷体"),
        body_font_size_pt=_median_or_default(all_body_sizes, 14.0),
        body_bold=False,
        section_heading_font_name=_most_common(all_body_fonts, "楷体"),
        section_heading_font_size_pt=_median_or_default(all_body_sizes, 14.0),
        section_heading_bold=True,
        doc_heading_font_name=_most_common(all_heading_fonts, "楷体"),
        doc_heading_bold=False,
        doc_heading_char_spacing=True,
    )

    # ─── Aggregate paragraph format ─────────
    all_indents = []
    all_line_spacings = []
    all_alignments = []

    for p in all_para:
        all_indents.extend(p["indents"])
        all_line_spacings.extend(p["line_spacings"])
        all_alignments.extend(p["alignments"])

    paragraph_format = ParagraphFormat(
        body_first_line_indent_cm=_median_or_default(all_indents, 0.99),
        body_line_spacing_emu=int(_median_or_default(
            [float(x) for x in all_line_spacings], 292100.0
        )),
        body_alignment=_most_common(all_alignments, "JUSTIFY (3)").split(" ")[0]
        if all_alignments else "JUSTIFY",
        section_heading_alignment="JUSTIFY",
        section_heading_indent_cm=0.0,
        doc_heading_alignment="CENTER",
    )

    # ─── Aggregate document structure ─────────
    # All patents should have 5 sections
    num_sections_list = [s["num_sections"] for s in all_structure]
    all_desc_sections = []
    for s in all_structure:
        if s["description_sections"]:
            all_desc_sections.append(s["description_sections"])

    # Find the most common description section sequence
    desc_section_order = ["技术领域", "背景技术", "发明内容", "附图说明", "具体实施方式"]
    if all_desc_sections:
        # Use the most common sequence
        seq_counter = Counter(tuple(s) for s in all_desc_sections)
        most_common_seq = seq_counter.most_common(1)[0][0]
        desc_section_order = list(most_common_seq)

    document_structure = DocumentStructure(
        section_order=["说明书摘要", "摘要附图", "权利要求书", "说明书", "说明书附图"],
        description_sections=desc_section_order,
        num_sections=int(_median_or_default([float(x) for x in num_sections_list], 5.0)),
        heading_char_separator="    ",
    )

    # ─── Aggregate content patterns ─────────
    all_claim_counts = [c["num_claims"] for c in all_content]
    all_indep_counts = [c["num_independent_claims"] for c in all_content]
    all_claim_lens = []
    for c in all_content:
        all_claim_lens.extend(c["claim_lengths"])
    all_abstract_lens = [c["abstract_length"] for c in all_content if c["abstract_length"] > 0]
    all_desc_lens = [c["description_length"] for c in all_content if c["description_length"] > 0]

    # Check pattern presence across all docs
    has_char_marker = sum(1 for c in all_content if c["has_characteristic_marker"])
    has_dep_format = sum(1 for c in all_content if c["has_dependent_format"])
    has_abstract_prefix = sum(1 for c in all_content if c["abstract_starts_with_benfaming"])
    has_explanation = sum(1 for c in all_content if c["has_explanation_markers"])
    has_closing = sum(1 for c in all_content if c["has_closing_statement"])

    content_patterns = ContentPatterns(
        claim_independent_prefix="其特征在于" if has_char_marker > len(all_content) // 2 else "",
        claim_dependent_format="根据权利要求{n}所述的" if has_dep_format > len(all_content) // 2 else "",
        claim_numbering_separator="、",
        tech_field_opening="本发明涉及",
        background_transition="然而",
        invention_summary_opening="为解决上述技术问题",
        detailed_desc_opening="为使本发明实施例的目的、技术方案和优点更加清楚" if has_explanation > len(all_content) // 2 else "",
        detailed_desc_explanation_marker="需要说明的是" if has_explanation > len(all_content) // 2 else "",
        detailed_desc_closing="最后应说明的是" if has_closing > len(all_content) // 2 else "",
        abstract_opening="本发明提供一种" if has_abstract_prefix > len(all_content) // 2 else "本发明",
        avg_claims_count=round(statistics.mean(all_claim_counts), 1) if all_claim_counts else 10.0,
        avg_independent_claims=round(statistics.mean(all_indep_counts), 1) if all_indep_counts else 1.0,
        avg_claim_length_chars=round(statistics.mean(all_claim_lens), 1) if all_claim_lens else 80.0,
        avg_description_total_chars=round(statistics.mean(all_desc_lens), 0) if all_desc_lens else 10000.0,
        avg_abstract_length_chars=round(statistics.mean(all_abstract_lens), 0) if all_abstract_lens else 250.0,
    )

    profile = PatentFeatureProfile(
        version="1.0",
        source_dir=str(docx_dir),
        num_patents_analyzed=len(b_files),
        extraction_date=datetime.now().isoformat(),
        page_layout=page_layout,
        typography=typography,
        paragraph_format=paragraph_format,
        document_structure=document_structure,
        content_patterns=content_patterns,
    )

    logger.info(
        f"特征提取完成: {len(b_files)} 篇专利 → "
        f"字体={typography.body_font_name} {typography.body_font_size_pt}pt, "
        f"缩进={paragraph_format.body_first_line_indent_cm}cm, "
        f"行距={paragraph_format.body_line_spacing_emu}EMU"
    )

    return profile


def extract_and_persist(
    docx_dir: Path, output_path: Optional[Path] = None
) -> PatentFeatureProfile:
    """Extract features from docx directory and save to JSON."""
    profile = extract_profile(docx_dir)
    target = output_path or DEFAULT_PROFILE_PATH
    profile.save(target)
    logger.info(f"特征 profile 已保存: {target}")
    return profile
