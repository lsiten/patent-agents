"""
Patent DOCX Generator Tool — 生成符合专利局规范的专利申请文件(.docx)

自包含工具，不依赖 document_gen 模块。
格式规范来自 PatentFeatureProfile（内置信定稿文件格式特征）：
- 楷体14pt正文，首行缩进0.99cm，左对齐
- A4页面，正确的各节页边距
- 文档标题字符间距展开（"权    利    要    求    书"）
- 说明书子标题加粗（技术领域、背景技术等）
- 权利要求编号格式："1、...其特征在于..."

用法（作为 Hermes Agent Tool）：
  agent 调用 → 传入结构化专利内容 → 生成规范 .docx → 返回文件路径

也支持直接命令行调用：
  python patent_docx_generator.py --title "..." --claims-json '...' --desc-json '...' --abstract "..." --figures-dir ./figures
"""

import json
import os
import re
import subprocess
import sys
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Dict, List, Optional, Set, Tuple

from docx.shared import Pt
from loguru import logger


# ═══════════════════════════════════════════════════════════════════
# PatentFeatureProfile — 内联格式描述
# ═══════════════════════════════════════════════════════════════════

@dataclass
class _PageLayout:
    page_width_cm: float = 21.0
    page_height_cm: float = 29.7
    margins_abstract: Tuple[float, float, float, float] = (2.60, 2.00, 2.80, 1.80)
    margins_abstract_drawings: Tuple[float, float, float, float] = (2.60, 2.00, 2.60, 2.00)
    margins_claims: Tuple[float, float, float, float] = (2.60, 2.00, 2.60, 2.00)
    margins_description: Tuple[float, float, float, float] = (2.60, 2.00, 2.80, 1.80)
    margins_description_drawings: Tuple[float, float, float, float] = (2.60, 2.00, 2.80, 1.80)


@dataclass
class _Typography:
    body_font_name: str = "楷体"
    body_font_size_pt: float = 14.0
    body_bold: bool = False
    section_heading_font_name: str = "楷体"
    section_heading_font_size_pt: float = 14.0
    section_heading_bold: bool = True
    doc_heading_font_name: str = "楷体"
    doc_heading_bold: bool = False
    doc_heading_char_spacing: bool = True


@dataclass
class _ParagraphFormat:
    body_first_line_indent_cm: float = 0.99
    body_line_spacing_emu: int = 292100
    body_alignment: str = "LEFT"
    section_heading_alignment: str = "LEFT"
    section_heading_indent_cm: float = 0.0
    doc_heading_alignment: str = "CENTER"


@dataclass
class _DocumentStructure:
    section_order: List[str] = field(default_factory=lambda: [
        "说明书摘要", "摘要附图", "权利要求书", "说明书", "说明书附图",
    ])
    description_sections: List[str] = field(default_factory=lambda: [
        "技术领域", "背景技术", "发明内容", "附图说明", "具体实施方式",
    ])
    num_sections: int = 5
    heading_char_separator: str = "    "


@dataclass
class _ContentPatterns:
    claim_independent_prefix: str = "其特征在于"
    claim_dependent_format: str = "根据权利要求{n}所述的"
    claim_numbering_separator: str = "、"


@dataclass
class _Profile:
    """Inline feature profile — no external file dependency."""
    version: str = "1.0"
    page_layout: _PageLayout = field(default_factory=_PageLayout)
    typography: _Typography = field(default_factory=_Typography)
    paragraph_format: _ParagraphFormat = field(default_factory=_ParagraphFormat)
    document_structure: _DocumentStructure = field(default_factory=_DocumentStructure)
    content_patterns: _ContentPatterns = field(default_factory=_ContentPatterns)


_profile: Optional[_Profile] = None


def _get_profile() -> _Profile:
    global _profile
    if _profile is None:
        _profile = _Profile()
    return _profile


# ═══════════════════════════════════════════════════════════════════
# 底层 DOCX 辅助函数
# ═══════════════════════════════════════════════════════════════════

def _set_line_spacing_exact(para, emu: int) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    twips = round(emu / 635)
    pPr = para._element.get_or_add_pPr()
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        pPr.append(spacing)
    spacing.set(qn("w:line"), str(twips))
    spacing.set(qn("w:lineRule"), "exact")


def _set_run_font(run, font_name: str, size_pt: float, bold: bool = False):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)


def _get_alignment(align_str: str):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    mapping = {
        "JUSTIFY": WD_ALIGN_PARAGRAPH.JUSTIFY,
        "CENTER": WD_ALIGN_PARAGRAPH.CENTER,
        "LEFT": WD_ALIGN_PARAGRAPH.LEFT,
        "RIGHT": WD_ALIGN_PARAGRAPH.RIGHT,
    }
    return mapping.get(align_str.upper(), WD_ALIGN_PARAGRAPH.LEFT)


def add_body_paragraph(
    doc,
    text: str,
    profile: _Profile,
    first_line_indent: bool = True,
    bold: bool = False,
) -> None:
    """Add a body text paragraph formatted according to the profile."""
    from docx.shared import Cm
    para = doc.add_paragraph()
    run = para.add_run(text)
    typo = profile.typography
    pfmt = profile.paragraph_format
    _set_run_font(run, typo.body_font_name, typo.body_font_size_pt, bold)
    para.paragraph_format.alignment = _get_alignment(pfmt.body_alignment)
    if first_line_indent:
        para.paragraph_format.first_line_indent = Cm(pfmt.body_first_line_indent_cm)
    _set_line_spacing_exact(para, pfmt.body_line_spacing_emu)


def add_section_heading(doc, text: str, profile: _Profile) -> None:
    from docx.shared import Cm
    para = doc.add_paragraph()
    run = para.add_run(text)
    typo = profile.typography
    pfmt = profile.paragraph_format
    _set_run_font(run, typo.section_heading_font_name, typo.section_heading_font_size_pt, bold=typo.section_heading_bold)
    para.paragraph_format.alignment = _get_alignment(pfmt.section_heading_alignment)
    _set_line_spacing_exact(para, pfmt.body_line_spacing_emu)


def add_document_heading(doc, text: str, profile: _Profile) -> None:
    from docx.shared import Cm
    struct = profile.document_structure
    pfmt = profile.paragraph_format
    typo = profile.typography
    spaced_text = struct.heading_char_separator.join(list(text))
    h = doc.add_heading(spaced_text, level=1)
    for run in h.runs:
        _set_run_font(run, typo.doc_heading_font_name, typo.section_heading_font_size_pt, bold=typo.doc_heading_bold)
    h.alignment = _get_alignment(pfmt.doc_heading_alignment)


def _set_section_margins(section, margins: Tuple[float, float, float, float]) -> None:
    from docx.shared import Cm
    top, bottom, left, right = margins
    section.top_margin = Cm(top)
    section.bottom_margin = Cm(bottom)
    section.left_margin = Cm(left)
    section.right_margin = Cm(right)


def _set_page_size(section, profile: _Profile) -> None:
    from docx.shared import Cm
    section.page_width = Cm(profile.page_layout.page_width_cm)
    section.page_height = Cm(profile.page_layout.page_height_cm)


def _get_margins_for_section(section_name: str, profile: _Profile) -> Tuple[float, float, float, float]:
    layout = profile.page_layout
    mapping = {
        "摘要": layout.margins_abstract,
        "摘要附图": layout.margins_abstract_drawings,
        "权利要求书": layout.margins_claims,
        "说明书": layout.margins_description,
        "说明书附图": layout.margins_description_drawings,
    }
    return mapping.get(section_name, layout.margins_abstract)


def add_new_section(doc, section_name: str, profile: _Profile) -> None:
    doc.add_section()
    new_section = doc.sections[-1]
    _set_page_size(new_section, profile)
    margins = _get_margins_for_section(section_name, profile)
    _set_section_margins(new_section, margins)


def _add_multiline_content(doc, content: Any, profile: _Profile) -> None:
    """Add content that may contain multiple paragraphs. Strips Markdown syntax."""
    content = _strip_markdown(content)
    paragraphs = [p.strip() for p in content.split("\n") if p.strip()]
    for para_text in paragraphs:
        add_body_paragraph(doc, para_text, profile)


def _add_figure_picture(doc, fig_info: Dict[str, str], profile: _Profile, width_inches: float = 5.0) -> bool:
    """Add a drawing image and caption to the current DOCX position."""
    try:
        from docx.shared import Inches as _Inches

        figure_number = _coerce_text(fig_info.get("figure_number")) or ""
        title = _coerce_text(fig_info.get("title")) or "专利附图"
        caption = f"{figure_number} {title}".strip()
        add_body_paragraph(doc, caption, profile, first_line_indent=False, bold=True)
        doc.add_picture(fig_info["path"], width=_Inches(width_inches))
        doc.add_paragraph("")
        return True
    except Exception as e:
        logger.warning(
            f"Failed to add figure {fig_info.get('figure_number')}: {e} "
            f"(path: {fig_info.get('path', '')})"
        )
        return False


def _add_multiline_content_with_figures(
    doc,
    content: Any,
    profile: _Profile,
    figures_by_number: Dict[str, Dict[str, str]],
    inserted_figures: Set[str],
) -> None:
    """Add paragraphs and insert the first referenced drawing immediately after its paragraph."""
    content = _strip_markdown(content)
    paragraphs = [p.strip() for p in content.split("\n") if p.strip()]
    for para_text in paragraphs:
        add_body_paragraph(doc, para_text, profile)
        for figure_number in sorted(figures_by_number.keys(), key=lambda value: int(re.search(r"\d+", value).group(0)) if re.search(r"\d+", value) else 0):
            if figure_number in inserted_figures:
                continue
            if re.search(rf"(?<!\d){re.escape(figure_number)}(?!\d)", para_text):
                if _add_figure_picture(doc, figures_by_number[figure_number], profile):
                    inserted_figures.add(figure_number)


def _coerce_text(value: Any) -> str:
    """Convert structured agent/tool output fields into document text."""
    if value is None:
        return ""
    if isinstance(value, str):
        return value
    if isinstance(value, dict):
        for key in ("content", "text", "value", "raw_response", "summary"):
            nested = value.get(key)
            if nested:
                return _coerce_text(nested)

        parts = []
        for item in value.values():
            text = _coerce_text(item).strip()
            if text:
                parts.append(text)
        return "\n".join(parts)
    if isinstance(value, list):
        parts = []
        for item in value:
            text = _coerce_text(item).strip()
            if text:
                parts.append(text)
        return "\n".join(parts)
    return str(value)


def _strip_markdown(text: Any) -> str:
    """Remove Markdown formatting from LLM-generated text for clean docx output."""
    text = _coerce_text(text)
    if not text:
        return text
    lines = text.split("\n")
    cleaned_lines = []
    for line in lines:
        # Skip pure horizontal rules
        if re.match(r'^[\s]*[-*_]{3,}\s*$', line):
            continue
        # Remove heading markers (# ## ### etc.)
        line = re.sub(r'^#{1,6}\s+', '', line)
        # Remove bold markers (**text** or __text__)
        line = re.sub(r'\*\*(.+?)\*\*', r'\1', line)
        line = re.sub(r'__(.+?)__', r'\1', line)
        # Remove italic markers (*text* or _text_)
        line = re.sub(r'(?<!\w)\*([^*]+?)\*(?!\w)', r'\1', line)
        line = re.sub(r'(?<!\w)_([^_]+?)_(?!\w)', r'\1', line)
        # Remove inline code backticks
        line = re.sub(r'`([^`]+?)`', r'\1', line)
        # Remove code block markers
        if re.match(r'^```', line):
            continue
        # Convert list markers to proper paragraph text
        line = re.sub(r'^\s*[-*+]\s+', '', line)
        # Remove link syntax [text](url)
        line = re.sub(r'\[([^\]]+?)\]\([^)]+?\)', r'\1', line)
        # Remove image syntax
        line = re.sub(r'!\[([^\]]*?)\]\([^)]+?\)', r'\1', line)
        # Remove blockquote markers
        line = re.sub(r'^>\s*', '', line)
        cleaned_lines.append(line)
    return "\n".join(cleaned_lines)


# ═══════════════════════════════════════════════════════════════════
# 专利附图生成 — 委托到外部脚本 scripts/generate_patent_figures.py
# ═══════════════════════════════════════════════════════════════════

def _generate_patent_figures(
    tech_description: str,
    task_id: str,
    output_dir: Optional[Path] = None,
) -> List[Dict[str, str]]:
    """
    生成专利附图。委托到 scripts/generate_patent_figures.py。

    支持两种策略（由外部脚本处理）：
    1. matplotlib 绘制系统架构图/流程图（默认）
    2. gen-img (gpt-image-2) AI 生成（设 IMAGE_GEN_* / LLM_* 环境变量）

    Returns:
        [{"path": "绝对路径", "title": "...", "figure_number": int}]
    """
    # 附图生成到对应专利 task 的 export 目录下
    backend_dir = Path(__file__).resolve().parent.parent.parent.parent.parent
    if output_dir is None:
        output_dir = backend_dir / "exports" / (task_id or "default") / "figures"
    output_dir.mkdir(parents=True, exist_ok=True)

    scripts_dir = backend_dir.parent / "scripts"
    script_path = scripts_dir / "generate_patent_figures.py"
    if not script_path.exists():
        logger.warning(f"Patent figures script not found: {script_path}")
        return []

    cmd = [
        sys.executable, str(script_path),
        "--tech-desc", tech_description[:2000],
        "--output-dir", str(output_dir),
        "--json",
    ]
    # 检测 AI 配置：IMAGE_GEN_* > LLM_*
    has_img_config = any(
        k.startswith("IMAGE_GEN_") and k.endswith("_API_KEY")
        for k in os.environ
    )
    has_llm_config = any(
        k.startswith("LLM_") and k.endswith("_API_KEY")
        for k in os.environ
    )
    if has_img_config or has_llm_config:
        cmd.append("--ai")

    try:
        result = subprocess.run(
            cmd, capture_output=True, text=True, timeout=180,
        )
        if result.returncode == 0:
            marker = "--- JSON ---"
            idx = result.stdout.find(marker)
            if idx >= 0:
                json_str = result.stdout[idx + len(marker):].strip()
                figures = json.loads(json_str)
                # 确保所有路径为绝对路径
                for f in figures:
                    fig_path = Path(f["path"])
                    if not fig_path.is_absolute():
                        # 相对路径基于 output_dir 解析
                        resolved = output_dir / fig_path.name
                        if resolved.exists():
                            f["path"] = str(resolved)
                        else:
                            # 尝试基于 backend_dir 解析
                            resolved = backend_dir / fig_path
                            f["path"] = str(resolved)
                    logger.info(f"Figure {f.get('figure_number')}: {f.get('title')} → {f.get('path')}")
                return figures
        else:
            logger.warning(f"Figure script failed (rc={result.returncode}): {result.stderr[:500]}")
    except Exception as e:
        logger.warning(f"Figure generation failed: {e}")

    return []


def _normalize_provided_figures(drawings: List[Dict[str, Any]]) -> List[Dict[str, str]]:
    """Convert patent_drawing_generator metadata into DOCX-ready figure entries."""
    figures: List[Dict[str, str]] = []
    backend_dir = Path(__file__).resolve().parent.parent.parent.parent.parent

    for index, drawing in enumerate(drawings, 1):
        if not isinstance(drawing, dict):
            continue
        path_value = drawing.get("file_path") or drawing.get("path")
        if not path_value:
            continue
        path = Path(str(path_value))
        if not path.is_absolute():
            path = backend_dir / path
        if not path.is_file():
            logger.warning(f"Provided patent drawing does not exist: {path}")
            continue
        if path.suffix.lower() not in {".png", ".jpg", ".jpeg", ".bmp", ".gif", ".tiff"}:
            logger.warning(f"Provided patent drawing is not directly supported by python-docx: {path}")
            continue
        figures.append(
            {
                "path": str(path),
                "title": _coerce_text(drawing.get("title")) or f"专利附图{index}",
                "figure_number": _coerce_text(drawing.get("figure_number")) or f"图{index}",
                "description": _coerce_text(drawing.get("description")),
            }
        )

    return figures


# ═══════════════════════════════════════════════════════════════════
# PDF/A 简化处理
# ═══════════════════════════════════════════════════════════════════

def _convert_to_pdfa(docx_path: str) -> str:
    """
    尝试将 DOCX 转换为 PDF/A（需要 LibreOffice）。
    如果不成功，返回原始 DOCX 路径。
    """
    try:
        pdf_path = docx_path.replace(".docx", ".pdf")
        subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "pdf", docx_path,
             "--outdir", str(Path(docx_path).parent)],
            capture_output=True, text=True, timeout=60,
        )
        if os.path.exists(pdf_path):
            logger.info(f"PDF/A 已生成: {pdf_path}")
            return pdf_path
    except Exception as e:
        logger.warning(f"PDF 转换失败（可忽略）: {e}")
    return docx_path


# ═══════════════════════════════════════════════════════════════════
# 主生成逻辑（Hermes Agent Tool 入口）
# ═══════════════════════════════════════════════════════════════════

class PatentDocxGeneratorTool:
    """生成专利申请DOCX文件的工具"""

    name = "patent_docx_generator"
    description = "将结构化的专利撰写结果生成为符合专利局规范的.docx文件"

    async def execute(
        self,
        title: Any = "专利申请文件",
        claims: Optional[Dict[str, Any]] = None,
        description: Optional[Dict[str, Any]] = None,
        abstract: Any = "",
        task_id: str = "",
        tech_description: str = "",
        drawings: Optional[List[Dict[str, Any]]] = None,
        **kwargs: Any,
    ) -> Dict[str, Any]:
        """
        生成专利DOCX文件（含附图）。

        Args:
            title: 专利标题
            claims: 权利要求 {"independent_claim": "...", "dependent_claims": ["..."]}
            description: 说明书 {"technical_field": "...", "background_art": "...",
                          "summary_of_invention": "...", "detailed_description": "...",
                          "description_of_drawings": "..."}
            abstract: 说明书摘要
            task_id: 任务ID
            tech_description: 原始技术方案描述（仅兼容旧调用；优先使用 drawings）
            drawings: patent_drawing_generator 已生成的附图元数据列表

        Returns:
            {"success": True, "file_path": "...", "figures": [...], "message": "..."}
        """
        if not claims and not description and not abstract:
            return {"success": False, "error": "未提供任何专利内容，无法生成文件"}

        from docx import Document
        from docx.shared import Pt

        claims = claims or {}
        description = description or {}

        try:
            profile = _get_profile()
            doc = Document()

            # ── 首页配置 ──
            first_section = doc.sections[0]
            _set_page_size(first_section, profile)
            margins = _get_margins_for_section("摘要", profile)
            _set_section_margins(first_section, margins)

            title = _coerce_text(title) or "专利申请文件"
            abstract = _coerce_text(abstract)

            # ── 说明书摘要 ── (仅在有内容时生成)
            if abstract and abstract.strip():
                add_document_heading(doc, "说明书摘要", profile)
                add_body_paragraph(doc, _strip_markdown(abstract), profile)

            figure_paths = _normalize_provided_figures(drawings or [])
            if not figure_paths and tech_description:
                figure_paths = _generate_patent_figures(tech_description, task_id)
            figures_by_number = {
                fig["figure_number"]: fig
                for fig in figure_paths
                if fig.get("figure_number") and fig.get("path")
            }
            inline_inserted_figures: Set[str] = set()

            # ── 摘要附图 ── (仅在有附图时生成)
            if figure_paths:
                add_new_section(doc, "摘要附图", profile)
                add_document_heading(doc, "摘要附图", profile)
                try:
                    from docx.shared import Inches as _Inches
                    doc.add_picture(figure_paths[0]["path"], width=_Inches(5.0))
                except Exception as e:
                    logger.warning(f"Failed to add abstract figure: {e} (path: {figure_paths[0].get('path', '')})")

            # ── 权利要求书 ── (仅在有内容时生成)
            if claims:
                add_new_section(doc, "权利要求书", profile)
                add_document_heading(doc, "权利要求书", profile)

                ind_claim = _strip_markdown(claims.get("independent_claim", ""))
                if ind_claim:
                    ind_claim = re.sub(r'^\d+[\.\、]\s*', '', ind_claim.strip())
                    add_body_paragraph(doc, f"1、{ind_claim}", profile)

                for i, dep in enumerate(claims.get("dependent_claims", []), 2):
                    dep_text = _strip_markdown(dep)
                    dep_text = re.sub(r'^\d+[\.\、]\s*', '', dep_text.strip())
                    add_body_paragraph(doc, f"{i}、{dep_text}", profile)

            # ── 说明书 ── (仅在有 description 内容时生成)
            has_description_content = any([
                description.get("technical_field"),
                description.get("background_art"),
                description.get("summary_of_invention"),
                description.get("description_of_drawings"),
                description.get("detailed_description"),
            ])
            
            if has_description_content:
                add_new_section(doc, "说明书", profile)
                add_document_heading(doc, "说明书", profile)

                # 专利名称（16pt楷体）
                title_para = doc.add_paragraph()
                title_run = title_para.add_run(_strip_markdown(title))
                _set_run_font(title_run, "楷体", 16.0)

                # 技术领域 (仅在有内容时生成)
                tech_field = description.get("technical_field", "")
                if isinstance(tech_field, dict):
                    tech_field = tech_field.get("content", "") or str(tech_field)
                if tech_field and tech_field.strip():
                    add_section_heading(doc, "技术领域", profile)
                    _add_multiline_content(doc, tech_field, profile)

                # 背景技术 (仅在有内容时生成)
                background = description.get("background_art", "")
                if isinstance(background, dict):
                    background = background.get("content", "") or str(background)
                if background and background.strip():
                    add_section_heading(doc, "背景技术", profile)
                    _add_multiline_content(doc, background, profile)

                # 发明内容 (仅在有内容时生成)
                summary = description.get("summary_of_invention", "")
                if isinstance(summary, dict):
                    # 结构化格式：包含 technical_problem, technical_solution, beneficial_effects
                    parts = []
                    if summary.get("technical_problem"):
                        parts.append(f"本发明要解决的技术问题是：{summary['technical_problem']}")
                    if summary.get("technical_solution"):
                        parts.append(f"\n{summary['technical_solution']}")
                    if summary.get("beneficial_effects"):
                        effects = summary['beneficial_effects']
                        if isinstance(effects, str):
                            parts.append(f"\n本发明的有益效果包括：{effects}")
                        elif isinstance(effects, list):
                            parts.append(f"\n本发明的有益效果包括：" + "；".join(str(e) for e in effects))
                    summary = "\n".join(parts) if parts else str(summary)
                if summary and summary.strip():
                    add_section_heading(doc, "发明内容", profile)
                    _add_multiline_content_with_figures(
                        doc,
                        summary,
                        profile,
                        figures_by_number,
                        inline_inserted_figures,
                    )

                # 附图说明（仅在有内容时生成）
                drawings_desc = description.get("description_of_drawings") or description.get("drawings_description", "")
                if isinstance(drawings_desc, dict):
                    drawings_desc = drawings_desc.get("content", "") or str(drawings_desc)
                if drawings_desc and drawings_desc.strip():
                    add_section_heading(doc, "附图说明", profile)
                    _add_multiline_content_with_figures(
                        doc,
                        drawings_desc,
                        profile,
                        figures_by_number,
                        inline_inserted_figures,
                    )

            # 具体实施方式 (仅在有内容时生成)
            detailed = description.get("detailed_description", "")
            if isinstance(detailed, dict):
                detailed = detailed.get("content", "") or str(detailed)
            elif isinstance(detailed, list):
                # 可能是实施例列表
                parts = []
                for item in detailed:
                    if isinstance(item, dict):
                        title_str = item.get("title", "")
                        content_str = item.get("content", "")
                        if title_str:
                            parts.append(f"{title_str}\n{content_str}")
                        else:
                            parts.append(content_str)
                    else:
                        parts.append(str(item))
                detailed = "\n\n".join(parts)
            if detailed and detailed.strip():
                add_section_heading(doc, "具体实施方式", profile)
                _add_multiline_content_with_figures(
                    doc,
                    detailed,
                    profile,
                    figures_by_number,
                    inline_inserted_figures,
                )

            # ── 说明书附图 ── (仅在有附图时生成)
            if figure_paths:
                add_new_section(doc, "说明书附图", profile)
                add_document_heading(doc, "说明书附图", profile)
                from docx.shared import Inches as _Inches
                for fig_info in figure_paths:
                    try:
                        fig_title = f"{fig_info.get('figure_number', '')}: {fig_info.get('title', '')}"
                        add_body_paragraph(doc, fig_title, profile, first_line_indent=False, bold=True)
                        doc.add_picture(fig_info["path"], width=_Inches(5.0))
                        doc.add_paragraph("")
                    except Exception as e:
                        logger.warning(f"Failed to add figure {fig_info.get('figure_number')}: {e} (path: {fig_info.get('path', '')})")

            # ── 保存文件 ──
            backend_dir = Path(__file__).resolve().parent.parent.parent.parent.parent
            export_dir = backend_dir / "exports" / (task_id or "default") / "final"
            export_dir.mkdir(parents=True, exist_ok=True)

            safe_title = re.sub(r'[\\/:*?"<>|]', '', title)[:50] or "专利申请书"
            file_path = export_dir / f"{safe_title}.docx"
            doc.save(str(file_path))

            result = {
                "success": True,
                "file_path": str(file_path),
                "file_name": f"{safe_title}.docx",
                "figures": figure_paths,
                "message": f"专利申请文件已生成：{file_path}",
                "sections": ["说明书摘要", "摘要附图", "权利要求书", "说明书", "说明书附图"],
            }

            # 尝试 PDF 转换（可选）
            if os.environ.get("ENABLE_PDF_EXPORT"):
                pdf_path = _convert_to_pdfa(str(file_path))
                result["pdf_path"] = pdf_path

            return result

        except Exception as e:
            logger.exception("DOCX generation failed")
            return {
                "success": False,
                "error": str(e),
                "message": f"生成DOCX文件失败：{str(e)}",
                "figures": [],
            }


# ═══════════════════════════════════════════════════════════════════
# CLI 入口（用于独立测试）
# ═══════════════════════════════════════════════════════════════════

def main():
    """CLI: python patent_docx_generator.py --title "..." --claims-json '...' ..."""
    import argparse
    parser = argparse.ArgumentParser(description="Generate patent DOCX")
    parser.add_argument("--title", default="专利申请文件")
    parser.add_argument("--claims-json", default="{}")
    parser.add_argument("--desc-json", default="{}")
    parser.add_argument("--abstract", default="")
    parser.add_argument("--task-id", default="cli_test")
    parser.add_argument("--tech-description", default="")
    parser.add_argument("--figures-dir", default="")
    args = parser.parse_args()

    import asyncio

    async def _run():
        tool = PatentDocxGeneratorTool()
        result = await tool.execute(
            title=args.title,
            claims=json.loads(args.claims_json),
            description=json.loads(args.desc_json),
            abstract=args.abstract,
            task_id=args.task_id,
            tech_description=args.tech_description,
        )
        print(json.dumps(result, ensure_ascii=False, indent=2))

    asyncio.run(_run())


if __name__ == "__main__":
    main()
